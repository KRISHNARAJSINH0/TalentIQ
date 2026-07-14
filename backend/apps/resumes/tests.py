import io
import zipfile
import json
from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.urls import reverse
from rest_framework import status
from rest_framework.test import APITestCase

from .models import Resume, ResumeSection, ConfidenceScore, SemanticValidation, ResumeError, RecoveryLog, ConsistencyReport, FieldSource, FieldHistory, SelfHealingReport, CopilotConversation, CopilotAction

User = get_user_model()


class ResumeAPITestCase(APITestCase):
    def setUp(self):
        # Create users
        self.user1 = User.objects.create_user(
            username="testuser1",
            email="test1@example.com",
            password="Password123!",
            first_name="Test",
            last_name="User"
        )
        self.user2 = User.objects.create_user(
            username="testuser2",
            email="test2@example.com",
            password="Password123!",
            first_name="Other",
            last_name="User"
        )

        # URLs
        self.upload_url = reverse("resumes:resume-upload")
        self.list_url = reverse("resumes:resume-list")
        self.history_url = reverse("resumes:resume-history")

        # Helpers for creating mock files
        # 1. Valid PDF file (starts with %PDF, meets 1KB min size)
        self.valid_pdf_content = b"%PDF-1.4\n" + (b"a" * 1500)
        self.valid_pdf = SimpleUploadedFile(
            "resume.pdf",
            self.valid_pdf_content,
            content_type="application/pdf"
        )

        # 2. Valid DOCX file (must be a valid zip, meets 1KB min size)
        docx_buffer = io.BytesIO()
        with zipfile.ZipFile(docx_buffer, "w") as zf:
            zf.writestr("word/document.xml", "<w:document xmlns:w='...'><w:body></w:body></w:document>")
            # Pad it to make it over 1KB
            zf.writestr("dummy.txt", "a" * 1500)
        self.valid_docx_content = docx_buffer.getvalue()
        self.valid_docx = SimpleUploadedFile(
            "resume.docx",
            self.valid_docx_content,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # 3. Invalid extension/MIME
        self.invalid_file = SimpleUploadedFile(
            "resume.txt",
            b"a" * 2000,
            content_type="text/plain"
        )

        # 4. Large file (exceeds 10MB)
        self.large_file = SimpleUploadedFile(
            "large.pdf",
            b"%PDF-1.4\n" + (b"a" * (11 * 1024 * 1024)),
            content_type="application/pdf"
        )

        # 5. Corrupted PDF (missing %PDF- header)
        self.corrupted_pdf = SimpleUploadedFile(
            "corrupt.pdf",
            b"not_a_pdf" + (b"a" * 1500),
            content_type="application/pdf"
        )

        # 6. Corrupted DOCX (not a valid zip file)
        self.corrupted_docx = SimpleUploadedFile(
            "corrupt.docx",
            b"not_a_zip" + (b"a" * 1500),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    def test_unauthenticated_access_blocked(self):
        """Verify anonymous requests to resumes API are rejected."""
        response = self.client.get(self.list_url)
        self.assertEqual(response.status_code, status.HTTP_401_UNAUTHORIZED)

    def test_successful_pdf_upload(self):
        """Verify valid PDF upload works, populates details, and sets is_active."""
        self.client.force_authenticate(user=self.user1)
        
        response = self.client.post(
            self.upload_url,
            {"original_file": self.valid_pdf, "resume_title": "My PDF Resume"},
            format="multipart"
        )
        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        self.assertEqual(response.data["filename"], "resume.pdf")
        self.assertEqual(response.data["version"], 1)
        self.assertEqual(response.data["status"], "completed")

        # Verify DB entry
        resume = Resume.objects.get(id=response.data["id"])
        self.assertEqual(resume.resume_title, "My PDF Resume")
        self.assertEqual(resume.file_size, len(self.valid_pdf_content))
        self.assertEqual(resume.mime_type, "application/pdf")
        self.assertTrue(resume.is_active)

    def test_successful_docx_upload(self):
        """Verify valid DOCX upload works."""
        self.client.force_authenticate(user=self.user1)

        response = self.client.post(
            self.upload_url,
            {"original_file": self.valid_docx},
            format="multipart"
        )
        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        # Should fallback to filename without ext as title
        self.assertEqual(response.data["version"], 1)

        resume = Resume.objects.get(id=response.data["id"])
        self.assertEqual(resume.resume_title, "resume")
        self.assertEqual(resume.mime_type, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    def test_multiple_uploads_increments_version(self):
        """Verify uploading another file increments version and swaps active status."""
        self.client.force_authenticate(user=self.user1)

        # Upload 1
        r1 = self.client.post(self.upload_url, {"original_file": self.valid_pdf}, format="multipart")
        self.assertEqual(r1.data["version"], 1)
        
        # Re-create file for upload 2 to avoid any test client stream reuse issues
        valid_pdf2 = SimpleUploadedFile(
            "resume2.pdf",
            self.valid_pdf_content,
            content_type="application/pdf"
        )
        
        # Upload 2
        r2 = self.client.post(self.upload_url, {"original_file": valid_pdf2}, format="multipart")
        self.assertEqual(r2.data["version"], 2)

        # Verify in DB
        resume1 = Resume.objects.get(id=r1.data["id"])
        resume2 = Resume.objects.get(id=r2.data["id"])
        
        self.assertFalse(resume1.is_active)
        self.assertTrue(resume2.is_active)

    def test_large_file_rejection(self):
        """Verify files over 10MB are rejected."""
        self.client.force_authenticate(user=self.user1)
        response = self.client.post(
            self.upload_url,
            {"original_file": self.large_file},
            format="multipart"
        )
        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn("File is too large", str(response.data))

    def test_invalid_extension_rejection(self):
        """Verify TXT or other unsupported formats are rejected."""
        self.client.force_authenticate(user=self.user1)
        response = self.client.post(
            self.upload_url,
            {"original_file": self.invalid_file},
            format="multipart"
        )
        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn("Unsupported file extension", str(response.data))

    def test_corrupted_pdf_rejection(self):
        """Verify a PDF lacking proper magic bytes is rejected."""
        self.client.force_authenticate(user=self.user1)
        response = self.client.post(
            self.upload_url,
            {"original_file": self.corrupted_pdf},
            format="multipart"
        )
        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn("corrupted or is not a valid PDF", str(response.data))

    def test_corrupted_docx_rejection(self):
        """Verify a docx file that is not a valid zip archive is rejected."""
        self.client.force_authenticate(user=self.user1)
        response = self.client.post(
            self.upload_url,
            {"original_file": self.corrupted_docx},
            format="multipart"
        )
        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn("corrupted or is not a valid DOCX", str(response.data))

    def test_secure_download_authorization(self):
        """Verify users can only download their own resumes."""
        # Create a resume for User 1
        self.client.force_authenticate(user=self.user1)
        upload_resp = self.client.post(self.upload_url, {"original_file": self.valid_pdf}, format="multipart")
        resume_id = upload_resp.data["id"]

        download_url = reverse("resumes:resume-download", args=[resume_id])

        # Attempt download as User 2 (Should be blocked)
        self.client.force_authenticate(user=self.user2)
        response = self.client.get(download_url)
        self.assertEqual(response.status_code, status.HTTP_403_FORBIDDEN)

        # Attempt download as User 1 (Should succeed)
        self.client.force_authenticate(user=self.user1)
        response = self.client.get(download_url)
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertTrue(response.streaming)
        self.assertEqual(response.headers["Content-Type"], "application/pdf")

    def test_activate_version(self):
        """Verify users can switch which resume version is active."""
        self.client.force_authenticate(user=self.user1)
        
        r1 = self.client.post(self.upload_url, {"original_file": self.valid_pdf}, format="multipart")
        
        valid_pdf2 = SimpleUploadedFile(
            "resume2.pdf",
            self.valid_pdf_content,
            content_type="application/pdf"
        )
        r2 = self.client.post(self.upload_url, {"original_file": valid_pdf2}, format="multipart")
        
        # Currently, r2 is active. Activate r1 manually.
        activate_url = reverse("resumes:resume-activate", args=[r1.data["id"]])
        response = self.client.patch(activate_url)
        self.assertEqual(response.status_code, status.HTTP_200_OK)

        # Verify DB
        resume1 = Resume.objects.get(id=r1.data["id"])
        resume2 = Resume.objects.get(id=r2.data["id"])
        self.assertTrue(resume1.is_active)
        self.assertFalse(resume2.is_active)

    def test_soft_delete(self):
        """Verify deleting a resume soft deletes it (sets is_deleted=True, doesn't delete file)."""
        self.client.force_authenticate(user=self.user1)
        upload_resp = self.client.post(self.upload_url, {"original_file": self.valid_pdf}, format="multipart")
        resume_id = upload_resp.data["id"]

        detail_url = reverse("resumes:resume-detail", args=[resume_id])

        # Delete
        delete_resp = self.client.delete(detail_url)
        self.assertEqual(delete_resp.status_code, status.HTTP_200_OK)

        # Verify it cannot be retrieved via API
        get_resp = self.client.get(detail_url)
        self.assertEqual(get_resp.status_code, status.HTTP_404_NOT_FOUND)

        # Verify it still exists in the database with is_deleted=True
        resume = Resume.all_objects.get(id=resume_id)
        self.assertTrue(resume.is_deleted)

    def test_pdf_extraction_success(self):
        """Verify clean text extraction from a valid PDF."""
        # pyrefly: ignore [missing-import]
        import fitz
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text(fitz.Point(50, 50), "Hello world from PyMuPDF test document.")
        # Insert enough text to cross the 1 KB minimum size threshold
        for i in range(20):
            page.insert_text(fitz.Point(50, 70 + i*15), f"This is line {i} of a dummy text sequence designed to increase the output size of this PDF to be safely above 1KB.")
        pdf_bytes = doc.write()
        doc.close()

        valid_pdf_file = SimpleUploadedFile("test.pdf", pdf_bytes, content_type="application/pdf")
        self.client.force_authenticate(user=self.user1)
        
        upload_resp = self.client.post(self.upload_url, {"original_file": valid_pdf_file}, format="multipart")
        self.assertEqual(upload_resp.status_code, status.HTTP_201_CREATED)
        resume_id = upload_resp.data["id"]

        # Trigger extraction
        extract_url = reverse("resumes:resume-extract", args=[resume_id])
        response = self.client.post(extract_url)
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(response.data["status"], "completed")
        self.assertGreater(response.data["text_length"], 0)

        # Fetch status
        status_url = reverse("resumes:resume-status-info", args=[resume_id])
        status_resp = self.client.get(status_url)
        self.assertEqual(status_resp.status_code, status.HTTP_200_OK)
        self.assertEqual(status_resp.data["status"], "completed")

        # Fetch text
        text_url = reverse("resumes:resume-text", args=[resume_id])
        text_resp = self.client.get(text_url)
        self.assertEqual(text_resp.status_code, status.HTTP_200_OK)
        self.assertIn("Hello world from PyMuPDF test document.", text_resp.data["extracted_text"])

    def test_docx_extraction_success(self):
        """Verify text extraction from a valid DOCX."""
        # pyrefly: ignore [missing-import]
        import docx
        doc = docx.Document()
        doc.add_paragraph("Paragraph text here.")
        doc.add_paragraph("Bullet item", style="List Bullet")
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "CellA"
        table.cell(0, 1).text = "CellB"

        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_bytes = docx_buffer.getvalue()

        valid_docx_file = SimpleUploadedFile(
            "test.docx", 
            docx_bytes, 
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        self.client.force_authenticate(user=self.user1)

        upload_resp = self.client.post(self.upload_url, {"original_file": valid_docx_file}, format="multipart")
        self.assertEqual(upload_resp.status_code, status.HTTP_201_CREATED)
        resume_id = upload_resp.data["id"]

        # Trigger extraction
        extract_url = reverse("resumes:resume-extract", args=[resume_id])
        response = self.client.post(extract_url)
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(response.data["status"], "completed")

        # Fetch text
        text_url = reverse("resumes:resume-text", args=[resume_id])
        text_resp = self.client.get(text_url)
        self.assertEqual(text_resp.status_code, status.HTTP_200_OK)
        self.assertIn("Paragraph text here.", text_resp.data["extracted_text"])
        self.assertIn("Bullet item", text_resp.data["extracted_text"])
        self.assertIn("CellA | CellB", text_resp.data["extracted_text"])

    def test_extraction_security_gate(self):
        """Verify users cannot trigger extraction or view text of other users' resumes."""
        self.client.force_authenticate(user=self.user1)
        upload_resp = self.client.post(self.upload_url, {"original_file": self.valid_pdf}, format="multipart")
        resume_id = upload_resp.data["id"]

        extract_url = reverse("resumes:resume-extract", args=[resume_id])
        text_url = reverse("resumes:resume-text", args=[resume_id])
        status_url = reverse("resumes:resume-status-info", args=[resume_id])

        # Access as user 2 (Should be blocked)
        self.client.force_authenticate(user=self.user2)
        
        response1 = self.client.post(extract_url)
        self.assertEqual(response1.status_code, status.HTTP_404_NOT_FOUND)

        response2 = self.client.get(text_url)
        self.assertEqual(response2.status_code, status.HTTP_404_NOT_FOUND)

        response3 = self.client.get(status_url)
        self.assertEqual(response3.status_code, status.HTTP_404_NOT_FOUND)


class ResumeRegexAPITestCase(APITestCase):
    def setUp(self):
        # Create users
        self.user1 = User.objects.create_user(
            username="regexuser1",
            email="regex1@example.com",
            password="Password123!",
            first_name="Regex",
            last_name="Tester"
        )
        self.user2 = User.objects.create_user(
            username="regexuser2",
            email="regex2@example.com",
            password="Password123!",
            first_name="Other",
            last_name="Tester"
        )

        # Create resume for user1 with raw extracted text
        self.resume = Resume.objects.create(
            user=self.user1,
            resume_title="Regex Test Resume",
            original_file=SimpleUploadedFile("dummy.pdf", b"pdf content", content_type="application/pdf"),
            original_filename="resume.pdf",
            stored_filename="resume_secure_123.pdf",
            file_size=1000,
            mime_type="application/pdf",
            extracted_text=(
                "John Doe\n"
                "Email: john.doe@gmail.com or primary@yahoo.com or john.doe@gmail.com\n"
                "Phone: +91 98765-43210 and +1 (302) 555-0123 and 9876543210\n"
                "Links:\n"
                "linkedin.com/in/johndoe\n"
                "https://github.com/johndoe\n"
                "www.johndoe.dev\n"
                "http://stackoverflow.com/users/johndoe\n"
                "https://twitter.com/johndoe\n"
                "Address: 123 tech park road, flat 4B, Bangalore.\n"
                "Pincode: 560001\n"
            ),
            extraction_status=Resume.ExtractionStatus.COMPLETED
        )

        self.regex_url = reverse("resumes:resume-regex", args=[self.resume.id])
        self.status_url = reverse("resumes:resume-regex-status", args=[self.resume.id])

    def test_regex_extraction_success(self):
        """Verify successful regex analysis, normalization, and JSON output."""
        self.client.force_authenticate(user=self.user1)

        # Trigger analysis
        response = self.client.post(self.regex_url)
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(response.data["status"], "completed")

        # Verify details
        data = response.data["regex_json"]
        self.assertEqual(data["email"], "john.doe@gmail.com")
        self.assertIn("primary@yahoo.com", data["secondary_emails"])
        # Duplicates removed (john.doe@gmail.com only appears once as primary)
        self.assertNotIn("john.doe@gmail.com", data["secondary_emails"])

        # Phone numbers: +919876543210, +13025550123.
        self.assertEqual(data["phone"], "+919876543210")
        self.assertIn("+13025550123", data["secondary_phones"])

        # URLs
        self.assertEqual(data["linkedin"], "https://linkedin.com/in/johndoe")
        self.assertEqual(data["github"], "https://github.com/johndoe")
        self.assertEqual(data["portfolio"], "https://www.johndoe.dev")
        self.assertEqual(data["stackoverflow"], "http://stackoverflow.com/users/johndoe")
        self.assertEqual(data["twitter"], "https://twitter.com/johndoe")

        # Address & Zip
        self.assertEqual(data["pincode"], "560001")
        self.assertIn("123 tech park road", data["address"])

        # Fetch status
        status_resp = self.client.get(self.status_url)
        self.assertEqual(status_resp.status_code, status.HTTP_200_OK)
        self.assertEqual(status_resp.data["status"], "completed")

        # Fetch JSON via GET
        get_resp = self.client.get(self.regex_url)
        self.assertEqual(get_resp.status_code, status.HTTP_200_OK)
        self.assertEqual(get_resp.data["regex_json"], data)

    def test_regex_security_gate(self):
        """Verify only the owner can run and access regex details."""
        # Access as user 2 (Should be blocked)
        self.client.force_authenticate(user=self.user2)

        post_resp = self.client.post(self.regex_url)
        self.assertEqual(post_resp.status_code, status.HTTP_404_NOT_FOUND)

        get_resp = self.client.get(self.regex_url)
        self.assertEqual(get_resp.status_code, status.HTTP_404_NOT_FOUND)

        status_resp = self.client.get(self.status_url)
        self.assertEqual(status_resp.status_code, status.HTTP_404_NOT_FOUND)


class ResumeSpacyAPITestCase(APITestCase):
    def setUp(self):
        # Create users
        self.user1 = User.objects.create_user(
            username="spacyuser1",
            email="spacy1@example.com",
            password="Password123!",
            first_name="Spacy",
            last_name="Tester"
        )
        self.user2 = User.objects.create_user(
            username="spacyuser2",
            email="spacy2@example.com",
            password="Password123!",
            first_name="Other",
            last_name="Tester"
        )

        # Create resume for user1 with raw extracted text containing clear entities
        self.resume = Resume.objects.create(
            user=self.user1,
            resume_title="Spacy Test Resume",
            original_file=SimpleUploadedFile("dummy.pdf", b"pdf content", content_type="application/pdf"),
            original_filename="resume.pdf",
            stored_filename="resume_secure_spacy_123.pdf",
            file_size=1000,
            mime_type="application/pdf",
            extracted_text=(
                "Alice Smith\n"
                "Experienced Software Engineer at Google.\n"
                "Worked from January 2020 to Present in Bangalore, India.\n"
                "Education: Bachelor of Technology from Stanford University."
            ),
            extraction_status=Resume.ExtractionStatus.COMPLETED
        )

        self.spacy_url = reverse("resumes:resume-spacy", args=[self.resume.id])
        self.status_url = reverse("resumes:resume-spacy-status", args=[self.resume.id])

    def test_spacy_extraction_success(self):
        """Verify successful spaCy NLP analysis, normalization, and JSON output structure."""
        self.client.force_authenticate(user=self.user1)

        # Trigger analysis
        response = self.client.post(self.spacy_url)
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(response.data["status"], "completed")

        # Verify details
        data = response.data["spacy_json"]
        self.assertIn("name", data)
        self.assertIn("organizations", data)
        self.assertIn("locations", data)
        self.assertIn("dates", data)
        self.assertIn("education_entities", data)
        self.assertIn("job_titles", data)

        # Check key expected entities
        # Bangalore/India should be in locations
        locations_lower = [l.lower() for l in data["locations"]]
        self.assertTrue(any("bangalore" in l or "india" in l for l in locations_lower))

        # Stanford University or Google should be in organizations
        orgs_lower = [o.lower() for o in data["organizations"]]
        self.assertTrue(any("google" in o or "stanford" in o for o in orgs_lower))

        # Fetch status
        status_resp = self.client.get(self.status_url)
        self.assertEqual(status_resp.status_code, status.HTTP_200_OK)
        self.assertEqual(status_resp.data["status"], "completed")

        # Fetch JSON via GET
        get_resp = self.client.get(self.spacy_url)
        self.assertEqual(get_resp.status_code, status.HTTP_200_OK)
        self.assertEqual(get_resp.data["spacy_json"], data)

    def test_spacy_security_gate(self):
        """Verify only the owner can run and access spaCy details."""
        # Access as user 2 (Should be blocked)
        self.client.force_authenticate(user=self.user2)

        post_resp = self.client.post(self.spacy_url)
        self.assertEqual(post_resp.status_code, status.HTTP_404_NOT_FOUND)

        get_resp = self.client.get(self.spacy_url)
        self.assertEqual(get_resp.status_code, status.HTTP_404_NOT_FOUND)

        status_resp = self.client.get(self.status_url)
        self.assertEqual(status_resp.status_code, status.HTTP_404_NOT_FOUND)


from unittest.mock import patch

class ResumeAIAPITestCase(APITestCase):
    def setUp(self):
        # Create users
        self.user1 = User.objects.create_user(
            username="aiuser1",
            email="ai1@example.com",
            password="Password123!",
            first_name="AI",
            last_name="Tester"
        )
        self.user2 = User.objects.create_user(
            username="aiuser2",
            email="ai2@example.com",
            password="Password123!",
            first_name="Other",
            last_name="Tester"
        )

        # Create resume for user1 with raw text, regex_json and spacy_json
        self.resume = Resume.objects.create(
            user=self.user1,
            resume_title="AI Test Resume",
            original_file=SimpleUploadedFile("dummy.pdf", b"pdf content", content_type="application/pdf"),
            original_filename="resume.pdf",
            stored_filename="resume_secure_ai_123.pdf",
            file_size=1000,
            mime_type="application/pdf",
            extracted_text=(
                "Dr. Jane Smith, MD\n"
                "Pediatrician with 5 years experience at Boston Children's Hospital.\n"
                "Education: Doctor of Medicine from Harvard Medical School."
            ),
            extraction_status=Resume.ExtractionStatus.COMPLETED,
            regex_status=Resume.ExtractionStatus.COMPLETED,
            regex_json={"email": "jane.smith@example.com", "phone": "+16175550199"},
            spacy_status=Resume.ExtractionStatus.COMPLETED,
            spacy_json={"name": "Jane Smith", "organizations": ["Boston Children's Hospital", "Harvard Medical School"]}
        )

        self.ai_url = reverse("resumes:resume-ai", args=[self.resume.id])
        self.status_url = reverse("resumes:resume-ai-status", args=[self.resume.id])

    @patch('apps.resumes.ai_service.GeminiService.generate_content')
    def test_ai_extraction_success(self, mock_generate_content):
        """Verify successful Gemini AI parsing and response structure."""
        self.client.force_authenticate(user=self.user1)

        # Setup mock response from GeminiService.generate_content
        mock_generate_content.return_value = json.dumps({
            "summary": "Pediatrician with 5 years experience at Boston Children's Hospital.",
            "job_role": "Pediatrician",
            "years_of_experience": "5 years",
            "current_company": "Boston Children's Hospital",
            "current_designation": "Pediatrician",
            "skills": ["Pediatrics", "Child Care"],
            "technical_skills": ["Pediatric Medicine", "Clinical Diagnosis"],
            "soft_skills": ["Communication", "Empathy"],
            "experience": [
                {
                    "company": "Boston Children's Hospital",
                    "designation": "Pediatrician",
                    "start_date": "2021",
                    "end_date": "Present",
                    "description": "Provide healthcare to children."
                }
            ],
            "education": [
                {
                    "institution": "Harvard Medical School",
                    "degree": "Doctor of Medicine",
                    "field_of_study": "Medicine",
                    "start_year": "2017",
                    "end_year": "2021"
                }
            ],
            "projects": [],
            "certifications": ["Medical License"],
            "awards": [],
            "achievements": [],
            "publications": [],
            "languages": ["English"],
            "hobbies": [],
            "references": []
        })

        # Trigger analysis
        response = self.client.post(self.ai_url)
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(response.data["status"], "completed")

        # Verify details
        data = response.data["ai_json"]
        self.assertEqual(data["job_role"], "Pediatrician")
        self.assertEqual(data["years_of_experience"], "5 years")
        self.assertEqual(data["current_company"], "Boston Children's Hospital")
        self.assertEqual(data["education"][0]["institution"], "Harvard Medical School")

        # Fetch status
        status_resp = self.client.get(self.status_url)
        self.assertEqual(status_resp.status_code, status.HTTP_200_OK)
        self.assertEqual(status_resp.data["status"], "completed")

        # Fetch JSON via GET
        get_resp = self.client.get(self.ai_url)
        self.assertEqual(get_resp.status_code, status.HTTP_200_OK)
        self.assertEqual(get_resp.data["ai_json"], data)

    def test_ai_security_gate(self):
        """Verify only the owner can run and access AI parsing details."""
        # Access as user 2 (Should be blocked)
        self.client.force_authenticate(user=self.user2)

        post_resp = self.client.post(self.ai_url)
        self.assertEqual(post_resp.status_code, status.HTTP_404_NOT_FOUND)

        get_resp = self.client.get(self.ai_url)
        self.assertEqual(get_resp.status_code, status.HTTP_404_NOT_FOUND)

        status_resp = self.client.get(self.status_url)
        self.assertEqual(status_resp.status_code, status.HTTP_404_NOT_FOUND)


class GeminiServiceTestCase(APITestCase):
    def test_prompt_builder(self):
        """Verify PromptBuilderService builds prompt correctly."""
        # pyrefly: ignore [missing-import]
        from apps.resumes.ai_service import PromptBuilderService
        prompt = PromptBuilderService.build_prompt(
            "John Doe resume",
            {"email": "john@example.com"},
            {"name": "John Doe"}
        )
        self.assertIn("John Doe resume", prompt)
        self.assertIn("john@example.com", prompt)
        self.assertIn("technical_skills", prompt)

    def test_gemini_service_initialization_with_missing_key(self):
        """Verify initialize_client sets use_mock to True when key is missing."""
        import os
        # pyrefly: ignore [missing-import]
        from apps.resumes.ai_service import GeminiService
        from django.conf import settings
        
        # Backup setting
        old_key = getattr(settings, "GEMINI_API_KEY", None)
        setattr(settings, "GEMINI_API_KEY", None)
        GeminiService._initialized = False
        
        # Remove from os.environ as well to test fallback
        old_env_key = os.environ.pop("GEMINI_API_KEY", None)
        
        try:
            GeminiService.initialize_client()
            self.assertTrue(GeminiService.use_mock)
        finally:
            # Restore setting and environment
            if old_key is not None:
                setattr(settings, "GEMINI_API_KEY", old_key)
            if old_env_key is not None:
                os.environ["GEMINI_API_KEY"] = old_env_key
            GeminiService._initialized = False


class MasterResumeTestCase(APITestCase):
    def setUp(self):
        # pyrefly: ignore [missing-import]
        from django.contrib.auth import get_user_model
        from apps.resumes.models import Resume
        
        User = get_user_model()
        self.user = User.objects.create_user(
            username="masteruser1",
            email="master1@example.com",
            password="Password123!"
        )
        self.other_user = User.objects.create_user(
            username="masteruser2",
            email="master2@example.com",
            password="Password123!"
        )
        
        self.resume = Resume.objects.create(
            user=self.user,
            resume_title="Test Resume",
            original_filename="test.pdf",
            stored_filename="test_stored.pdf",
            file_size=1024,
            mime_type="application/pdf",
            extraction_status="completed",
            regex_status="completed",
            spacy_status="completed",
            ai_status="completed",
            regex_json={
                "email": "regex@example.com",
                "phone": "+1 555-0199",
                "linkedin": "https://linkedin.com/in/regex",
                "github": "https://github.com/regex",
                "portfolio": "https://regex.me",
            },
            spacy_json={
                "name": "Jane Spacy",
                "organizations": ["TCS", "Infosys Ltd."],
                "dates": ["May 2021", "2022"],
            },
            ai_json={
                "name": "Jane Gemini",
                "email": "ai@example.com",
                "summary": "Experienced software developer.",
                "skills": ["JS", "Py", "Node", "reactjs", "Python"],
                "technical_skills": ["ReactJS", "JS"],
                "education": [
                    {
                        "institution": "Harvard University",
                        "degree": "B.S.",
                        "field_of_study": "Computer Science",
                        "start_year": "2018",
                        "end_year": "2022"
                    }
                ],
                "experience": [
                    {
                        "company": "TCS",
                        "designation": "Developer",
                        "start_date": "May 2021",
                        "end_date": "current",
                        "description": "Building stuff"
                    },
                    {
                        "company": "TCS",
                        "designation": "Developer",
                        "start_date": "May 2021",
                        "end_date": "current",
                        "description": "Building stuff"  # Duplicate experience entry
                    }
                ]
            }
        )
        
        self.merge_url = f"/api/resumes/{self.resume.id}/merge/"
        self.master_url = f"/api/resumes/{self.resume.id}/master/"
        self.completion_url = f"/api/resumes/{self.resume.id}/completion/"
        self.client.force_authenticate(user=self.user)

    def test_conflict_resolution_and_normalization(self):
        """Verify priority-based merging and normalizations work correctly."""
        # pyrefly: ignore [missing-import]
        from apps.resumes.validation_service import MasterResumeBuilder
        
        builder = MasterResumeBuilder()
        success = builder.build_master_profile(self.resume)
        self.assertTrue(success)
        
        self.resume.refresh_from_db()
        self.assertEqual(self.resume.validation_status, "completed")
        self.assertTrue(self.resume.completion_percentage > 0)
        
        master = self.resume.master_resume_json
        
        # Verify Priority
        # Regex Priority for contact
        self.assertEqual(master["email"], "regex@example.com")
        self.assertEqual(master["phone"], "+1 555-0199")
        self.assertEqual(master["linkedin"], "https://linkedin.com/in/regex")
        
        # spaCy Priority for Name
        self.assertEqual(master["name"], "Jane Spacy")
        
        # Gemini Priority for Summary
        self.assertEqual(master["summary"], "Experienced software developer.")
        
        # Verify Skill Normalization and Deduplication (JS -> JavaScript, Py -> Python, node -> Node.js, reactjs -> React, etc.)
        self.assertIn("JavaScript", master["skills"])
        self.assertIn("Python", master["skills"])
        self.assertIn("Node.js", master["skills"])
        self.assertIn("React", master["skills"])
        self.assertNotIn("JS", master["skills"])
        
        # Verify Company Normalization (TCS -> Tata Consultancy Services)
        self.assertEqual(master["experience"][0]["company"], "Tata Consultancy Services")
        
        # Verify Date Normalization ("current" -> "Present", "May 2021" -> "05/2021")
        self.assertEqual(master["experience"][0]["start_date"], "05/2021")
        self.assertEqual(master["experience"][0]["end_date"], "Present")
        
        # Verify Deduplication: only one experience entry (the second was a duplicate key)
        self.assertEqual(len(master["experience"]), 1)

    def test_completion_percentage_calculation(self):
        """Verify completion percentage returns correct scoring weights."""
        # pyrefly: ignore [missing-import]
        from apps.resumes.validation_service import MasterResumeBuilder
        builder = MasterResumeBuilder()
        
        # 1. Full Profile (should be high)
        full_profile = {
            "name": "Jane Doe",
            "email": "jane@example.com",
            "phone": "555-1234",
            "summary": "Summary",
            "skills": ["JavaScript"],
            "experience": [{"company": "A"}],
            "education": [{"institution": "B"}],
            "projects": [{"title": "C"}],
            "certifications": ["D"],
            "languages": ["English"],
        }
        score = builder.calculate_completion(full_profile)
        self.assertEqual(score, 100.0)
        
        # 2. Minimal Profile
        min_profile = {
            "name": "Jane Doe",
        }
        score_min = builder.calculate_completion(min_profile)
        self.assertEqual(score_min, 5.0)

    def test_master_profile_endpoints(self):
        """Verify merge, master profile retrieval, and completion endpoints."""
        # POST to merge
        response = self.client.post(self.merge_url)
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.data["validation_status"], "completed")
        self.assertTrue(response.data["completion_percentage"] > 0)
        
        # GET master
        response = self.client.get(self.master_url)
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.data["validation_status"], "completed")
        self.assertIn("name", response.data["master_resume_json"])
        
        # GET completion
        response = self.client.get(self.completion_url)
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.data["validation_status"], "completed")
        self.assertTrue(response.data["completion_percentage"] > 0)

    def test_master_profile_security(self):
        """Verify other users cannot access or run master profile features."""
        self.client.force_authenticate(user=self.other_user)
        
        response = self.client.post(self.merge_url)
        self.assertEqual(response.status_code, 404)
        
        response = self.client.get(self.master_url)
        self.assertEqual(response.status_code, 404)
        
        response = self.client.get(self.completion_url)
        self.assertEqual(response.status_code, 404)


class ResumeSectionDetectionTestCase(APITestCase):
    def setUp(self):
        self.user = User.objects.create_user(
            username="secuser1",
            email="sec1@example.com",
            password="Password123!"
        )
        self.other_user = User.objects.create_user(
            username="secuser2",
            email="sec2@example.com",
            password="Password123!"
        )

        self.resume = Resume.objects.create(
            user=self.user,
            resume_title="Section Testing Resume",
            original_filename="resume.pdf",
            file_size=1024,
            mime_type="application/pdf",
            extracted_text=(
                "John Doe\n"
                "Email: john@example.com\n\n"
                "SUMMARY\n"
                "Experienced full stack developer.\n\n"
                "WORK HISTORY\n"
                "Google - Software Engineer (2020 - Present)\n"
                "Built search services.\n\n"
                "EDUCATION\n"
                "Stanford University - B.S. in Computer Science\n"
            ),
            extraction_status=Resume.ExtractionStatus.COMPLETED
        )

        self.sections_url = reverse("resume-sections")

    def test_section_detection_by_resume_id(self):
        """Verify section detection by resume_id retrieves, parses, persists, and returns results."""
        self.client.force_authenticate(user=self.user)
        
        response = self.client.post(self.sections_url, {"resume_id": str(self.resume.id)}, format="json")
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        
        data = response.data
        self.assertIn("layout", data)
        self.assertEqual(data["layout"], "single_column")
        
        sections = data["sections"]
        self.assertEqual(len(sections), 4)
        
        # Check normalized types
        types = [s["type"] for s in sections]
        self.assertIn("summary", types)
        self.assertIn("experience", types)
        self.assertIn("education", types)
        
        # Verify persistence in DB
        db_sections = ResumeSection.objects.filter(resume=self.resume)
        self.assertEqual(db_sections.count(), 4)
        self.assertEqual(db_sections.filter(section_type="experience").first().title, "WORK HISTORY")

    def test_section_detection_by_raw_text(self):
        """Verify section detection by raw text returns layout and parsed sections without DB persistence."""
        self.client.force_authenticate(user=self.user)
        
        raw_text = (
            "Alice Smith\n"
            "Email: alice@example.com\n\n"
            "TECHNICAL SKILLS\n"
            "Python, Django, React, Postgres\n\n"
            "PROJECTS\n"
            "ResumeAI: building parsers\n"
        )
        
        response = self.client.post(self.sections_url, {"text": raw_text}, format="json")
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        
        data = response.data
        self.assertEqual(data["layout"], "single_column")
        self.assertEqual(len(data["sections"]), 3)
        
        types = [s["type"] for s in data["sections"]]
        self.assertIn("skills", types)
        self.assertIn("projects", types)

    def test_section_detection_security_gate(self):
        """Verify other users cannot run section detection on a resume."""
        self.client.force_authenticate(user=self.other_user)
        
        response = self.client.post(self.sections_url, {"resume_id": str(self.resume.id)}, format="json")
        self.assertEqual(response.status_code, status.HTTP_404_NOT_FOUND)

    def test_missing_parameters_returns_bad_request(self):
        """Verify request without parameters is rejected."""
        self.client.force_authenticate(user=self.user)
        
        response = self.client.post(self.sections_url, {}, format="json")
        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn("required", response.data["detail"])

    def test_layout_detection_heuristics(self):
        """Verify local heuristics detect non-standard column styles."""
        from apps.resumes.services.section_detector import LayoutDetector
        
        # Text with many separators -> table_layout
        table_text = "Name | Age | Role\nJohn Doe | 30 | Engineer\nJane Smith | 28 | Designer"
        self.assertEqual(LayoutDetector.detect_heuristically(table_text), "table_layout")
        
        # Standard text -> single_column
        std_text = "My name is John.\nI work at Google.\nI went to Stanford."
        self.assertEqual(LayoutDetector.detect_heuristically(std_text), "single_column")

    def test_title_normalization(self):
        """Verify SectionNormalizer correctly maps creative/non-standard headers."""
        from apps.resumes.services.section_detector import SectionNormalizer
        
        self.assertEqual(SectionNormalizer.normalize("WORK HISTORY"), "experience")
        self.assertEqual(SectionNormalizer.normalize("Academic Background"), "education")
        self.assertEqual(SectionNormalizer.normalize("Capabilities"), "skills")
        self.assertEqual(SectionNormalizer.normalize("Things I Built"), "projects")
        self.assertEqual(SectionNormalizer.normalize("Core competencies"), "skills")


class ResumeConfidenceTestCase(APITestCase):
    def setUp(self):
        self.user = User.objects.create_user(
            username="confuser1",
            email="conf1@example.com",
            password="Password123!"
        )
        self.other_user = User.objects.create_user(
            username="confuser2",
            email="conf2@example.com",
            password="Password123!"
        )

        # Create a mock resume with pre-filled parser jsons
        self.resume = Resume.objects.create(
            user=self.user,
            resume_title="Confidence Evaluation Resume",
            original_filename="cv.pdf",
            file_size=1024,
            mime_type="application/pdf",
            extracted_text=(
                "John Doe\n"
                "Email: john@example.com\n"
                "Phone: 555-9876\n\n"
                "EDUCATION\n"
                "Stanford University - Bachelor of Science\n\n"
                "TECHNICAL SKILLS\n"
                "Python, Django, React, Postgres\n"
            ),
            regex_json={
                "email": "john@example.com",
                "phone": "555-9876",
                "linkedin": "https://linkedin.com/in/johndoe"
            },
            spacy_json={
                "name": "John Doe",
                "email": "john@example.com"
            },
            ai_json={
                "name": "John Doe",
                "email": "john@example.com",
                "summary": "Experienced software developer.",
                "skills": ["Python", "Django", "React", "Postgres"],
                "education": [
                    {
                        "institution": "Stanford University",
                        "degree": "Bachelor of Science"
                    }
                ]
            },
            master_resume_json={
                "name": "John Doe",
                "email": "john@example.com",
                "phone": "555-9876",
                "linkedin": "https://linkedin.com/in/johndoe",
                "summary": "Experienced software developer.",
                "skills": ["Python", "Django", "React", "Postgres"],
                "education": [
                    {
                        "institution": "Stanford University",
                        "degree": "Bachelor of Science"
                    }
                ]
            },
            extraction_status=Resume.ExtractionStatus.COMPLETED
        )

        # Create sections
        ResumeSection.objects.create(
            resume=self.resume,
            section_type="skills",
            title="TECHNICAL SKILLS",
            content="Python, Django, React, Postgres",
            position=0
        )
        ResumeSection.objects.create(
            resume=self.resume,
            section_type="education",
            title="EDUCATION",
            content="Stanford University - Bachelor of Science",
            position=1
        )

        self.confidence_url = reverse("resume-confidence")
        self.detail_url = reverse("resume-confidence-detail", kwargs={"pk": self.resume.id})

    def test_confidence_calculations_logic(self):
        """Verify the confidence calculation rules (boosts, section matches, semantic drops)."""
        from apps.resumes.services.confidence_engine import ConfidenceEngine
        
        engine = ConfidenceEngine()
        res = engine.evaluate_resume(self.resume)
        
        # 1. Email: Regex source (99) + Agreement (100 override) or Email entity boost (+20) -> clamped to 100
        self.assertEqual(res["email"]["confidence"], 100.0)
        self.assertEqual(res["email"]["status"], "accepted")
        self.assertIn("Full agreement", res["email"]["reason"])
        
        # 2. Name: spaCy source (95) + PERSON entity boost (+10) + Header area boost (+10) -> clamped to 100
        self.assertEqual(res["name"]["confidence"], 100.0)
        self.assertEqual(res["name"]["status"], "accepted")
        
        # 3. Name Semantic check: Name containing "engineer" drops confidence
        self.resume.master_resume_json["name"] = "John Doe Software Engineer"
        self.resume.save()
        res_with_role = engine.evaluate_resume(self.resume)
        # 95 (spacy base) + 5 (agreement boost) + 10 (person boost) - 30 (role penalty) = 80.0
        self.assertEqual(res_with_role["name"]["confidence"], 80.0)
        self.assertIn("job title", res_with_role["name"]["reason"])

        # 4. Skills Semantic check: Skill containing "university" drops confidence
        self.resume.master_resume_json["skills"] = ["Python", "Stanford University"]
        self.resume.save()
        res_with_edu_skill = engine.evaluate_resume(self.resume)
        # Average of Python (90 + 10 skills boost = 100) and Stanford University (90 + 0 boost - 40 penalty = 50)
        # Average: (100 + 50) / 2 = 75.0
        self.assertEqual(res_with_edu_skill["skills"]["confidence"], 75.0)

    def test_api_calculate_and_retrieve(self):
        """Verify POST calculations and GET detail retrieve endpoints function and persist results."""
        self.client.force_authenticate(user=self.user)
        
        # POST to trigger calculations
        response = self.client.post(self.confidence_url, {"resume_id": str(self.resume.id)}, format="json")
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        
        data = response.data
        self.assertEqual(data["resume_id"], str(self.resume.id))
        self.assertIn("fields", data)
        self.assertIn("confidence_map", data)
        
        # Verify database storage
        db_scores = ConfidenceScore.objects.filter(resume=self.resume)
        self.assertTrue(db_scores.exists())
        self.assertEqual(db_scores.filter(field="email").first().confidence, 100.0)
        
        # GET to retrieve
        response_get = self.client.get(self.detail_url)
        self.assertEqual(response_get.status_code, status.HTTP_200_OK)
        self.assertEqual(response_get.data["resume_id"], str(self.resume.id))
        self.assertEqual(len(response_get.data["fields"]), db_scores.count())

    def test_confidence_security_gates(self):
        """Verify other users cannot run or view confidence scores for a resume."""
        self.client.force_authenticate(user=self.other_user)
        
        # POST block
        response = self.client.post(self.confidence_url, {"resume_id": str(self.resume.id)}, format="json")
        self.assertEqual(response.status_code, status.HTTP_404_NOT_FOUND)
        
        # GET block
        response_get = self.client.get(self.detail_url)
        self.assertEqual(response_get.status_code, status.HTTP_404_NOT_FOUND)


class SemanticValidatorEngineTestCase(APITestCase):
    """
    Test suite for KnowledgeBase, OntologyEngine, SemanticMatcher, EntityClassifier,
    SemanticValidator, and 11 distinct resume genres.
    """

    def setUp(self):
        self.user = User.objects.create_user(
            username="semantic_tester",
            email="semantic@example.com",
            password="Password123!"
        )
        self.client.force_authenticate(user=self.user)

        from apps.resumes.models import SemanticValidation
        from apps.resumes.services.knowledge_base import KnowledgeBase
        from apps.resumes.services.ontology_engine import OntologyEngine
        from apps.resumes.services.semantic_matcher import SemanticMatcher
        from apps.resumes.services.entity_classifier import EntityClassifier
        from apps.resumes.services.semantic_validator import SemanticValidator

        self.kb = KnowledgeBase()
        self.ontology = OntologyEngine()
        self.classifier = EntityClassifier()
        self.validator = SemanticValidator()

    def test_knowledge_base_loading(self):
        """Verify all 8 JSON dictionaries are loaded in KnowledgeBase."""
        self.assertTrue(self.kb.is_exact_entity("SKILL", "Python"))
        self.assertTrue(self.kb.is_exact_entity("COMPANY", "Google"))
        self.assertTrue(self.kb.is_exact_entity("UNIVERSITY", "MIT"))
        self.assertTrue(self.kb.is_exact_entity("DESIGNATION", "Software Engineer"))
        self.assertTrue(self.kb.is_exact_entity("CERTIFICATE", "AWS Certified Solutions Architect"))
        self.assertTrue(self.kb.is_exact_entity("LANGUAGE", "English"))

    def test_ontology_engine_rules(self):
        """Verify semantic ontology rules for person names, universities, dates, and designations."""
        # Person name with designation should fail
        is_valid, score, reason = self.ontology.validate_person_name("Software Engineer")
        self.assertFalse(is_valid)
        self.assertIn("designation", reason.lower())

        # Valid person name
        is_valid_name, score_name, _ = self.ontology.validate_person_name("John Doe")
        self.assertTrue(is_valid_name)
        self.assertGreaterEqual(score_name, 90.0)

        # University keyword matching
        is_univ, score_u, _ = self.ontology.validate_university("Massachusetts Institute of Technology")
        self.assertTrue(is_univ)
        self.assertGreaterEqual(score_u, 90.0)

        # Date format matching
        is_date, score_d, _ = self.ontology.validate_date("05/2024")
        self.assertTrue(is_date)
        self.assertGreaterEqual(score_d, 90.0)

    def test_semantic_matcher_similarity(self):
        """Verify vector cosine similarity and text matching."""
        from apps.resumes.services.semantic_matcher import SemanticMatcher
        sim = SemanticMatcher.cosine_similarity("Python Data Science", "Python Machine Learning")
        self.assertGreater(sim, 0.0)

        match_score = SemanticMatcher.match_against_category(
            "Stanford University",
            ["Stanford University", "MIT"],
            ["university", "college"]
        )
        self.assertEqual(match_score, 100.0)

    def test_entity_classifier(self):
        """Verify entity classification across validation categories."""
        res_univ = self.classifier.classify_entity("MIT")
        self.assertEqual(res_univ["top_category"], "UNIVERSITY")

        res_skill = self.classifier.classify_entity("React")
        self.assertIn(res_skill["top_category"], ["SKILL", "TECHNOLOGY"])

        res_desig = self.classifier.classify_entity("Backend Developer")
        self.assertIn(res_desig["top_category"], ["DESIGNATION", "ROLE"])

    def test_semantic_anomaly_detection_problem_example(self):
        """
        Tests the problem scenario from prompt specification:
        {
          "name": "Software Engineer",
          "education": ["Python"],
          "skills": ["MIT"],
          "company": "Backend Developer",
          "project": "Google"
        }
        """
        problem_payload = {
            "name": "Software Engineer",
            "education": ["Python"],
            "skills": ["MIT"],
            "company": "Backend Developer",
            "project": "Google"
        }

        report = self.validator.validate_payload(problem_payload)
        validations = {v["field"]: v for v in report["validations"]}

        # 1. MIT inside Skills -> Detected University, status invalid, action move_to_education
        mit_val = validations.get("skills[0]")
        self.assertIsNotNone(mit_val)
        self.assertEqual(mit_val["value"], "MIT")
        self.assertEqual(mit_val["category"], "University")
        self.assertEqual(mit_val["expected_category"], "Skill")
        self.assertEqual(mit_val["status"], "invalid")
        self.assertEqual(mit_val["action"], "move_to_education")
        self.assertIn("university", mit_val["reason"].lower())

        # 2. Python inside Education -> Detected Skill, status invalid, action move_to_skills
        py_val = validations.get("education[0]")
        self.assertIsNotNone(py_val)
        self.assertEqual(py_val["value"], "Python")
        self.assertIn(py_val["category"], ["Skill", "Technology"])
        self.assertEqual(py_val["status"], "invalid")
        self.assertEqual(py_val["action"], "move_to_skills")

        # 3. Software Engineer inside Name -> Detected Designation, status invalid
        name_val = validations.get("name")
        self.assertIsNotNone(name_val)
        self.assertEqual(name_val["value"], "Software Engineer")
        self.assertIn(name_val["category"], ["Designation", "Role"])
        self.assertEqual(name_val["status"], "invalid")

        # 4. Google inside Project -> Maybe company. Review.
        proj_val = validations.get("project")
        self.assertIsNotNone(proj_val)
        self.assertEqual(proj_val["value"], "Google")
        self.assertEqual(proj_val["action"], "review")
        self.assertEqual(proj_val["reason"], "Maybe company. Review.")

    def test_resume_genre_software(self):
        """Tests Software Engineer Resume."""
        payload = {
            "name": "Alex Mercer",
            "company": "Google",
            "education": ["Stanford University"],
            "skills": ["Python", "Docker", "Kubernetes", "Redis"],
            "certifications": ["AWS Certified Solutions Architect"]
        }
        report = self.validator.validate_payload(payload)
        self.assertGreaterEqual(report["metrics"]["semantic_accuracy"], 95.0)

    def test_resume_genre_academic_cv(self):
        """Tests Academic CV."""
        payload = {
            "name": "Dr. Eleanor Vance",
            "university": "University of Oxford",
            "publications": ["IEEE Journal of Quantum Computing 2024"],
            "awards": ["Gold Medalist Fellowship"]
        }
        report = self.validator.validate_payload(payload)
        self.assertGreaterEqual(report["metrics"]["semantic_accuracy"], 95.0)

    def test_resume_genre_designer(self):
        """Tests Designer Resume."""
        payload = {
            "name": "Sophia Chen",
            "designation": "UI/UX Designer",
            "skills": ["Figma", "Photoshop", "Wireframing"],
            "company": "Adobe"
        }
        report = self.validator.validate_payload(payload)
        self.assertGreaterEqual(report["metrics"]["semantic_accuracy"], 95.0)

    def test_resume_genre_medical_cv(self):
        """Tests Medical CV."""
        payload = {
            "name": "Dr. Robert House",
            "designation": "Physician",
            "university": "Harvard University",
            "skills": ["Communication", "Problem Solving"]
        }
        report = self.validator.validate_payload(payload)
        self.assertGreaterEqual(report["metrics"]["semantic_accuracy"], 95.0)

    def test_resume_genre_research_cv(self):
        """Tests Research CV."""
        payload = {
            "name": "Alan Turing",
            "university": "University of Cambridge",
            "skills": ["Machine Learning", "Mathematics"],
            "publications": ["ACM Computing Surveys"]
        }
        report = self.validator.validate_payload(payload)
        self.assertGreaterEqual(report["metrics"]["semantic_accuracy"], 95.0)

    def test_resume_genre_law(self):
        """Tests Law Resume."""
        payload = {
            "name": "Harvey Specter",
            "designation": "Attorney",
            "company": "Pearson Hardman",
            "university": "Harvard Law School"
        }
        report = self.validator.validate_payload(payload)
        self.assertGreaterEqual(report["metrics"]["semantic_accuracy"], 95.0)

    def test_resume_genre_teacher(self):
        """Tests Teacher Resume."""
        payload = {
            "name": "Clara Oswald",
            "designation": "High School Teacher",
            "university": "Delhi University",
            "languages": ["English", "French"]
        }
        report = self.validator.validate_payload(payload)
        self.assertGreaterEqual(report["metrics"]["semantic_accuracy"], 95.0)

    def test_resume_genre_student(self):
        """Tests Student Resume."""
        payload = {
            "name": "Rohan Sharma",
            "university": "BITS Pilani",
            "skills": ["Java", "C++", "SQL"],
            "projects": ["Portfolio Web App"]
        }
        report = self.validator.validate_payload(payload)
        self.assertGreaterEqual(report["metrics"]["semantic_accuracy"], 95.0)

    def test_resume_genre_two_column(self):
        """Tests Two-Column Resume."""
        payload = {
            "name": "David Miller",
            "designation": "DevOps Engineer",
            "skills": ["Terraform", "Ansible", "AWS"],
            "company": "Microsoft"
        }
        report = self.validator.validate_payload(payload)
        self.assertGreaterEqual(report["metrics"]["semantic_accuracy"], 95.0)

    def test_resume_genre_creative(self):
        """Tests Creative Resume."""
        payload = {
            "name": "Maya Lin",
            "designation": "Creative Director",
            "skills": ["Illustrator", "InDesign", "Prototyping"],
            "company": "Canva"
        }
        report = self.validator.validate_payload(payload)
        self.assertGreaterEqual(report["metrics"]["semantic_accuracy"], 95.0)

    def test_resume_genre_canva(self):
        """Tests Canva Resume."""
        payload = {
            "name": "Liam Nelson",
            "designation": "Product Manager",
            "company": "Salesforce",
            "skills": ["Jira", "Agile", "Scrum"]
        }
        report = self.validator.validate_payload(payload)
        self.assertGreaterEqual(report["metrics"]["semantic_accuracy"], 95.0)

    def test_semantic_validation_api_post(self):
        """Tests POST /api/resume/semantic/ endpoint with payload."""
        response = self.client.post("/api/resume/semantic/", {
            "payload": {
                "name": "John Doe",
                "skills": ["MIT"]
            }
        }, format="json")
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        data = response.json()
        self.assertIn("validations", data)
        self.assertIn("metrics", data)

    def test_semantic_validation_api_post_with_resume_id(self):
        """Tests POST /api/resume/semantic/ endpoint with resume_id."""
        resume = Resume.objects.create(
            user=self.user,
            resume_title="Test Resume",
            master_resume_json={"name": "Jane Smith", "skills": ["Python"]}
        )
        response = self.client.post("/api/resume/semantic/", {
            "resume_id": str(resume.id)
        }, format="json")
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        data = response.json()
        self.assertEqual(data["resume_id"], str(resume.id))
        self.assertTrue(SemanticValidation.objects.filter(resume=resume).exists())

    def test_semantic_validation_api_get_detail(self):
        """Tests GET /api/resume/semantic/{id} endpoint."""
        resume = Resume.objects.create(
            user=self.user,
            resume_title="Test Resume 2",
            master_resume_json={"name": "Jane Smith"}
        )
        val = SemanticValidation.objects.create(
            resume=resume,
            field="name",
            value="Jane Smith",
            category="Person",
            expected_category="Person",
            semantic_score=95.0,
            reason="Valid name",
            status="valid",
            action="accept"
        )
        response = self.client.get(f"/api/resume/semantic/{val.id}/")
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        data = response.json()
        self.assertEqual(data["field"], "name")
        self.assertEqual(data["category"], "Person")


class ErrorDetectorEngineTestCase(APITestCase):
    """
    Test suite for Stage 8 Error Detection Engine, covering:
    DuplicateDetector, TimelineValidator, ContactValidator, ConsistencyValidator,
    QualityValidator, SectionErrorDetector, ErrorDetector orchestrator, API endpoints,
    and 12 distinct resume genres.
    """

    def setUp(self):
        self.user = User.objects.create_user(
            username="errordetect_user",
            email="errordetect@example.com",
            password="Password123!"
        )
        self.client.force_authenticate(user=self.user)

        from apps.resumes.services.duplicate_detector import DuplicateDetector
        from apps.resumes.services.timeline_validator import TimelineValidator
        from apps.resumes.services.contact_validator import ContactValidator
        from apps.resumes.services.consistency_validator import ConsistencyValidator
        from apps.resumes.services.quality_validator import QualityValidator
        from apps.resumes.services.section_error_detector import SectionErrorDetector
        from apps.resumes.services.error_detector import ErrorDetector

        self.duplicate_detector = DuplicateDetector()
        self.timeline_validator = TimelineValidator()
        self.contact_validator = ContactValidator()
        self.consistency_validator = ConsistencyValidator()
        self.quality_validator = QualityValidator()
        self.section_error_detector = SectionErrorDetector()
        self.error_detector = ErrorDetector()

    def test_duplicate_detector_skills_and_companies(self):
        payload = {
            "skills": ["Python", "Django", "python", "Docker"],
            "experience": [
                {"company": "Google LLC", "start_year": 2020, "end_year": 2022},
                {"company": "Google Inc", "start_year": 2022, "end_year": 2024}
            ]
        }
        errors = self.duplicate_detector.detect_duplicates(payload)
        self.assertTrue(any(e["type"] == "duplicate_value" and "python" in e["value"].lower() for e in errors))
        self.assertTrue(any(e["type"] == "duplicate_value" and "Google Inc" in e["value"] for e in errors))

    def test_timeline_validator_negative_duration_and_overlaps(self):
        payload = {
            "experience": [
                {"company": "Company A", "start_date": "2023", "end_date": "2022"},
                {"company": "Company B", "start_date": "2020", "end_date": "present"},
                {"company": "Company C", "start_date": "2021", "end_date": "present"}
            ],
            "education": [
                {"school": "University X", "start_year": "2022", "end_year": "2018"}
            ]
        }
        errors = self.timeline_validator.validate_timelines(payload)
        self.assertTrue(any("Negative duration" in e["reason"] for e in errors))
        self.assertTrue(any("Multiple active current positions" in e["reason"] for e in errors))
        self.assertTrue(any("end year (2018) is before start year" in e["reason"] for e in errors))

    def test_contact_validator_missing_and_invalid(self):
        payload = {
            "email": "invalid_email_at_domain",
            "phone": "123",
            "linkedin": ""
        }
        errors = self.contact_validator.validate_contact(payload)
        self.assertTrue(any("Malformed email address" in e["reason"] for e in errors))
        self.assertTrue(any("Invalid phone number" in e["reason"] for e in errors))
        self.assertTrue(any("Missing LinkedIn profile" in e["reason"] for e in errors))

    def test_consistency_validator_student_and_fresher_contradictions(self):
        payload = {
            "name": "Student John",
            "designation": "Student Intern",
            "years_of_experience": 12,
            "summary": "Fresher looking for entry role",
            "current_designation": "Senior Lead Principal Engineer"
        }
        errors = self.consistency_validator.validate_consistency(payload)
        self.assertTrue(any(e["type"] == "contradictory_values" for e in errors))

    def test_section_error_detector_misplaced_entities(self):
        payload = {
            "name": "Software Engineer",
            "skills": ["MIT", "Python"],
            "education": ["Python Developer"],
            "projects": ["Google LLC"]
        }
        errors = self.section_error_detector.detect_section_errors(payload)
        self.assertTrue(any("Name field contains designation" in e["reason"] for e in errors))
        self.assertTrue(any("university" in e["reason"] for e in errors))

    def test_error_detector_full_orchestration(self):
        payload = {
            "name": "Alex Taylor",
            "email": "alex@example.com",
            "phone": "+1 555 019 2831",
            "skills": ["Python", "Python", "MIT"],
            "experience": [
                {"company": "Google", "start_date": "2023", "end_date": "2021"}
            ]
        }
        res = self.error_detector.detect_errors(payload)
        self.assertIn("errors", res)
        self.assertIn("metrics", res)
        self.assertGreaterEqual(res["metrics"]["total_errors"], 2)

    def test_error_detection_api_post_raw_payload(self):
        response = self.client.post("/api/resume/errors/", {
            "payload": {
                "skills": ["Python", "Python"],
                "email": ""
            }
        }, format="json")
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        data = response.json()
        self.assertIn("errors", data)

    def test_error_detection_api_post_with_resume_id(self):
        resume = Resume.objects.create(
            user=self.user,
            resume_title="Test Resume Errors",
            master_resume_json={"skills": ["Java", "Java"], "email": ""}
        )
        response = self.client.post("/api/resume/errors/", {
            "resume_id": str(resume.id)
        }, format="json")
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        data = response.json()
        self.assertEqual(data["resume_id"], str(resume.id))
        self.assertTrue(ResumeError.objects.filter(resume=resume).exists())

    def test_error_detection_api_summary(self):
        resume = Resume.objects.create(user=self.user, resume_title="Summary Test")
        ResumeError.objects.create(
            resume=resume,
            type="timeline_error",
            field="experience[0]",
            severity="critical",
            reason="Negative duration"
        )
        response = self.client.get("/api/resume/errors/summary/")
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        data = response.json()
        self.assertGreaterEqual(data["total_errors_detected"], 1)

    # 12 RESUME GENRES ERROR DETECTION COVERAGE
    def test_genre_1_software_engineer(self):
        payload = {"name": "Dev", "email": "dev@test.com", "phone": "1234567890", "skills": ["C++", "C++"], "experience": [{"company": "Meta", "start_year": 2020, "end_year": 2024}]}
        res = self.error_detector.detect_errors(payload)
        self.assertTrue(any(e["type"] == "duplicate_value" for e in res["errors"]))

    def test_genre_2_medical_cv(self):
        payload = {"name": "Dr. Sarah", "email": "sarah@med.org", "phone": "9876543210", "education": [{"school": "Harvard Medical", "start_year": 2024, "end_year": 2020}]}
        res = self.error_detector.detect_errors(payload)
        self.assertTrue(any("end year (2020) is before start year" in e["reason"] for e in res["errors"]))

    def test_genre_3_academic_cv(self):
        payload = {"name": "Prof. Smith", "email": "", "phone": "5551234567", "skills": ["Grant Writing"]}
        res = self.error_detector.detect_errors(payload)
        self.assertTrue(any(e["type"] == "missing_contact" for e in res["errors"]))

    def test_genre_4_designer_resume(self):
        payload = {"name": "UI Designer", "skills": ["Figma"], "email": "ui@design.io", "phone": "1234567890"}
        res = self.error_detector.detect_errors(payload)
        self.assertTrue(any(e["type"] == "wrong_entity" for e in res["errors"]))

    def test_genre_5_students(self):
        payload = {"designation": "Student", "years_of_experience": 15, "email": "st@edu.com", "phone": "1234567890"}
        res = self.error_detector.detect_errors(payload)
        self.assertTrue(any(e["type"] == "contradictory_values" for e in res["errors"]))

    def test_genre_6_freshers(self):
        payload = {"designation": "Fresher Senior Vice President", "email": "fr@test.com", "phone": "1234567890"}
        res = self.error_detector.detect_errors(payload)
        self.assertTrue(any(e["type"] == "contradictory_values" for e in res["errors"]))

    def test_genre_7_research_cv(self):
        payload = {"name": "Dr. Alice", "email": "alice@lab.org", "phone": "1234567890", "projects": ["Google", "Google"]}
        res = self.error_detector.detect_errors(payload)
        self.assertTrue(any(e["type"] == "duplicate_value" for e in res["errors"]))

    def test_genre_8_canva_resumes(self):
        payload = {"skills": ["Photoshop"], "email": "canva@user.com", "phone": "1234567890"}
        res = self.error_detector.detect_errors(payload)
        self.assertTrue(any(e["type"] == "missing_field" and e["field"] == "summary" for e in res["errors"]))

    def test_genre_9_multipage_resumes(self):
        payload = {"name": "Multi Page", "email": "multi@page.com", "phone": "1234567890", "skills": ["Python", "C++"], "education": [{"school": "Stanford"}], "experience": [{"company": "A", "start_year": 2010, "end_year": 2015}]}
        res = self.error_detector.detect_errors(payload)
        self.assertGreaterEqual(res["metrics"]["quality_score"], 60.0)

    def test_genre_10_twocolumn_resumes(self):
        payload = {"name": "Col User", "email": "col@test.com", "phone": "1234567890", "skills": ["Java"], "education": [{"school": "MIT"}]}
        res = self.error_detector.detect_errors(payload)
        self.assertIn("metrics", res)

    def test_genre_11_creative_layouts(self):
        payload = {"name": "Creative One", "email": "c@art.com", "phone": "1234567890", "certifications": ["AWS", "AWS"]}
        res = self.error_detector.detect_errors(payload)
        self.assertTrue(any(e["type"] == "duplicate_value" for e in res["errors"]))

    def test_genre_12_broken_resumes(self):
        payload = {"email": "broken_no_at_symbol", "phone": "123"}
        res = self.error_detector.detect_errors(payload)
        self.assertTrue(any(e["type"] in ["formatting_issue", "missing_field"] for e in res["errors"]))


class RecoveryEngineTestCase(APITestCase):
    """
    Test suite for Stage 9 AI Recovery Engine, covering:
    EntityMover, DuplicateResolver, DateRecovery, SummaryRecovery,
    RecoveryRules, RecoveryEngine orchestrator, API endpoints,
    and 13 distinct resume genres.
    """

    def setUp(self):
        self.user = User.objects.create_user(
            username="recovery_user",
            email="recovery@example.com",
            password="Password123!"
        )
        self.client.force_authenticate(user=self.user)

        from apps.resumes.services.entity_mover import EntityMover
        from apps.resumes.services.duplicate_resolver import DuplicateResolver
        from apps.resumes.services.date_recovery import DateRecovery
        from apps.resumes.services.summary_recovery import SummaryRecovery
        from apps.resumes.services.recovery_engine import RecoveryEngine

        self.entity_mover = EntityMover()
        self.duplicate_resolver = DuplicateResolver()
        self.date_recovery = DateRecovery()
        self.summary_recovery = SummaryRecovery()
        self.recovery_engine = RecoveryEngine()

    def test_entity_mover_designation_skills_companies(self):
        payload = {
            "name": "Software Engineer",
            "skills": ["MIT", "Python"],
            "education": ["Python Developer"],
            "projects": ["Google LLC"]
        }
        res = self.entity_mover.process_entity_movement(payload)
        rec_json = res["payload"]
        self.assertEqual(rec_json["designation"], "Software Engineer")
        self.assertIn("MIT", [e.get("institution") for e in rec_json["education"] if isinstance(e, dict)])
        self.assertIn("Google LLC", [e.get("company") for e in rec_json["experience"] if isinstance(e, dict)])

    def test_duplicate_resolver_case_folding_and_legal_suffixes(self):
        payload = {
            "skills": ["Python", "PYTHON", "python", "Docker"],
            "experience": [
                {"company": "Google Inc.", "start_year": 2020, "end_year": 2022},
                {"company": "Google LLC", "start_year": 2022, "end_year": 2024}
            ]
        }
        res = self.duplicate_resolver.resolve_duplicates(payload)
        rec_json = res["payload"]
        self.assertEqual(len(rec_json["skills"]), 2)
        self.assertEqual(rec_json["experience"][0]["company"], "Google")

    def test_date_recovery_inversion_and_normalization(self):
        payload = {
            "experience": [
                {"company": "Company A", "start_date": "2024", "end_date": "2022"},
                {"company": "Company B", "start_date": "2020", "end_date": "Current"}
            ]
        }
        res = self.date_recovery.recover_dates(payload)
        rec_json = res["payload"]
        self.assertEqual(rec_json["experience"][0]["start_date"], "2022")
        self.assertEqual(rec_json["experience"][0]["end_date"], "2024")
        self.assertEqual(rec_json["experience"][1]["end_date"], "Present")

    def test_summary_recovery_extraction_and_synthesis(self):
        payload = {
            "name": "Alex Smith",
            "skills": ["Java", "Spring Boot"],
            "experience": [
                {"company": "Tech Corp", "description": "Professional Summary: Accomplished developer with high performance."}
            ]
        }
        res = self.summary_recovery.recover_summary(payload)
        rec_json = res["payload"]
        self.assertIn("Accomplished developer", rec_json["summary"])

    def test_recovery_engine_full_flow(self):
        payload = {
            "name": "Software Engineer",
            "skills": ["MIT", "Python", "Python"],
            "education": ["Python Developer"],
            "experience": [
                {"company": "Google LLC", "start_date": "2024", "end_date": "2022"}
            ]
        }
        res = self.recovery_engine.recover_payload(payload)
        self.assertIn("recovered_json", res)
        self.assertIn("recoveries", res)
        self.assertIn("metrics", res)
        self.assertGreaterEqual(res["metrics"]["recovery_accuracy"], 95.0)

    def test_recovery_api_post_raw_payload(self):
        response = self.client.post("/api/recovery/", {
            "payload": {
                "name": "Software Engineer",
                "skills": ["MIT"]
            }
        }, format="json")
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        data = response.json()
        self.assertIn("recovered_json", data)

    def test_recovery_api_post_resume_id(self):
        resume = Resume.objects.create(
            user=self.user,
            resume_title="Test Recovery Resume",
            master_resume_json={"name": "Software Engineer", "skills": ["MIT"]}
        )
        response = self.client.post("/api/recovery/", {
            "resume_id": str(resume.id)
        }, format="json")
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        data = response.json()
        self.assertEqual(data["resume_id"], str(resume.id))
        self.assertTrue(RecoveryLog.objects.filter(resume=resume).exists())

    def test_recovery_api_history(self):
        resume = Resume.objects.create(user=self.user, resume_title="History Test")
        RecoveryLog.objects.create(
            resume=resume,
            field="education",
            previous_value="skills",
            new_value="MIT",
            confidence=97.0,
            status="recovered"
        )
        response = self.client.get("/api/recovery/history/")
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        data = response.json()
        self.assertGreaterEqual(data["history_count"], 1)

    # 13 RESUME GENRES RECOVERY COVERAGE
    def test_genre_1_software(self):
        res = self.recovery_engine.recover_payload({"skills": ["Java", "JAVA"], "experience": [{"company": "Amazon Inc"}]})
        self.assertEqual(res["recovered_json"]["experience"][0]["company"], "Amazon")

    def test_genre_2_academic(self):
        res = self.recovery_engine.recover_payload({"skills": ["Harvard University"], "email": "prof@edu.org"})
        self.assertTrue(any(r["to"] == "education" for r in res["recoveries"]))

    def test_genre_3_research_cv(self):
        res = self.recovery_engine.recover_payload({"projects": ["Meta Platforms Inc"]})
        self.assertTrue(any(r["to"] == "experience" for r in res["recoveries"]))

    def test_genre_4_designer(self):
        res = self.recovery_engine.recover_payload({"name": "UI UX Lead", "skills": ["Figma"]})
        self.assertEqual(res["recovered_json"]["designation"], "UI UX Lead")

    def test_genre_5_law(self):
        res = self.recovery_engine.recover_payload({"name": "Attorney At Law", "email": "attorney@law.com"})
        self.assertEqual(res["recovered_json"].get("designation") or res["recovered_json"].get("name"), "Attorney At Law")

    def test_genre_6_doctor(self):
        res = self.recovery_engine.recover_payload({"skills": ["Johns Hopkins School of Medicine"]})
        self.assertTrue(any(r["to"] == "education" for r in res["recoveries"]))

    def test_genre_7_teacher(self):
        res = self.recovery_engine.recover_payload({"experience": [{"company": "School District", "start_date": "2023", "end_date": "2021"}]})
        self.assertEqual(res["recovered_json"]["experience"][0]["start_date"], "2021")

    def test_genre_8_medical_cv(self):
        res = self.recovery_engine.recover_payload({"experience": [{"company": "General Hospital Corp", "end_date": "Current"}]})
        self.assertEqual(res["recovered_json"]["experience"][0]["end_date"], "Present")

    def test_genre_9_canva(self):
        res = self.recovery_engine.recover_payload({"skills": ["Photoshop", "Photoshop"]})
        self.assertEqual(len(res["recovered_json"]["skills"]), 1)

    def test_genre_10_twocolumn(self):
        res = self.recovery_engine.recover_payload({"education": ["Python Developer"]})
        self.assertTrue(any(r["to"] == "skills" for r in res["recoveries"]))

    def test_genre_11_multipage(self):
        res = self.recovery_engine.recover_payload({"name": "John Doe", "email": "john.doe@gmail.com"})
        self.assertGreaterEqual(res["metrics"]["recovery_accuracy"], 95.0)

    def test_genre_12_broken(self):
        res = self.recovery_engine.recover_payload({"email": "alex.taylor@tech.com"})
        self.assertEqual(res["recovered_json"]["name"], "Alex Taylor")

    def test_genre_13_malformed(self):
        res = self.recovery_engine.recover_payload({"skills": ["C++", "c++"]})
        self.assertEqual(len(res["recovered_json"]["skills"]), 1)


class ConsistencyCheckerTestCase(APITestCase):
    """
    Test suite for Stage 9 / Phase 9.6 Consistency Checker, covering:
    TimelineChecker, CareerChecker, RoleChecker, ProfileChecker,
    CompletenessChecker, ConsistencyChecker orchestrator, API endpoints,
    and 14 distinct professions.
    """

    def setUp(self):
        self.user = User.objects.create_user(
            username="consistency_user",
            email="consistency@example.com",
            password="Password123!"
        )
        self.client.force_authenticate(user=self.user)

        from apps.resumes.services.timeline_checker import TimelineChecker
        from apps.resumes.services.career_checker import CareerChecker
        from apps.resumes.services.role_checker import RoleChecker
        from apps.resumes.services.profile_checker import ProfileChecker
        from apps.resumes.services.completeness_checker import CompletenessChecker
        from apps.resumes.services.consistency_checker import ConsistencyChecker

        self.timeline_checker = TimelineChecker()
        self.career_checker = CareerChecker()
        self.role_checker = RoleChecker()
        self.profile_checker = ProfileChecker()
        self.completeness_checker = CompletenessChecker()
        self.consistency_checker = ConsistencyChecker()

    def test_timeline_checker_multiple_active_and_overlaps(self):
        payload = {
            "experience": [
                {"company": "Company A", "start_date": "2020", "end_date": "Present"},
                {"company": "Company B", "start_date": "2022", "end_date": "Present"}
            ]
        }
        issues = self.timeline_checker.check_timeline_consistency(payload)
        self.assertTrue(any("Multiple active current positions" in i["reason"] for i in issues))

    def test_career_checker_student_15_years_experience(self):
        payload = {
            "designation": "Student",
            "years_of_experience": 15
        }
        issues = self.career_checker.check_career_consistency(payload)
        self.assertTrue(any("Student" in i["reason"] for i in issues))

    def test_role_checker_backend_missing_skills(self):
        payload = {
            "designation": "Software Engineer",
            "skills": ["HTML", "CSS"]
        }
        issues, suggestions = self.role_checker.check_role_consistency(payload)
        self.assertTrue(any("Software Engineer" in i["reason"] for i in issues))
        self.assertTrue(len(suggestions) > 0)

    def test_profile_checker_backend_fashion_projects(self):
        payload = {
            "designation": "Backend Engineer",
            "projects": ["Fashion Catalog", "Restaurant Menu"]
        }
        issues = self.profile_checker.check_profile_consistency(payload)
        self.assertTrue(any("weak technical relevance" in i["reason"] for i in issues))

    def test_completeness_checker_scores(self):
        payload = {
            "name": "Alex Taylor",
            "email": "alex@example.com",
            "skills": ["Python"]
        }
        score, issues = self.completeness_checker.check_completeness(payload)
        self.assertGreater(score, 0)
        self.assertTrue(any("Missing Education" in i["reason"] for i in issues))

    def test_consistency_checker_orchestrator_full_flow(self):
        payload = {
            "name": "Jane Doe",
            "email": "jane@example.com",
            "phone": "555-1234",
            "designation": "Data Analyst",
            "skills": ["SQL", "Excel", "Power BI"],
            "education": [{"degree": "BS CS", "institution": "MIT"}],
            "experience": [{"company": "Tech Corp", "start_year": 2021, "end_year": "Present"}]
        }
        res = self.consistency_checker.check_consistency(payload)
        self.assertIn("consistency_score", res)
        self.assertIn("score_label", res)
        self.assertGreaterEqual(res["consistency_score"], 70)

    def test_consistency_api_post_raw_payload(self):
        response = self.client.post("/api/consistency/", {
            "payload": {
                "name": "Alex Taylor",
                "email": "alex@example.com",
                "skills": ["Python"]
            }
        }, format="json")
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        data = response.json()
        self.assertIn("consistency_score", data)

    def test_consistency_api_post_resume_id(self):
        resume = Resume.objects.create(
            user=self.user,
            resume_title="Test Consistency Resume",
            master_resume_json={"name": "Alex Taylor", "email": "alex@example.com", "skills": ["Python"]}
        )
        response = self.client.post("/api/consistency/", {
            "resume_id": str(resume.id)
        }, format="json")
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        data = response.json()
        self.assertEqual(str(data["resume"]), str(resume.id))
        self.assertTrue(ConsistencyReport.objects.filter(resume=resume).exists())

    def test_consistency_api_history(self):
        resume = Resume.objects.create(user=self.user, resume_title="History Test")
        ConsistencyReport.objects.create(
            resume=resume,
            score=92.0,
            score_label="Strong",
            issues=[]
        )
        response = self.client.get("/api/consistency/history/")
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        data = response.json()
        self.assertGreaterEqual(data["history_count"], 1)

    # 14 PROFESSIONS COVERAGE
    def test_profession_1_student(self):
        res = self.consistency_checker.check_consistency({"designation": "Student", "years_of_experience": 12})
        self.assertTrue(any("Student" in i["reason"] for i in res["issues"]))

    def test_profession_2_fresher(self):
        res = self.consistency_checker.check_consistency({"designation": "Fresher", "years_of_experience": 6})
        self.assertTrue(any(i["type"] == "career" for i in res["issues"]))

    def test_profession_3_software_engineer(self):
        res = self.consistency_checker.check_consistency({"designation": "Software Engineer", "skills": ["HTML", "CSS"]})
        self.assertTrue(any(i["type"] == "role_skills" for i in res["issues"]))

    def test_profession_4_researcher(self):
        res = self.consistency_checker.check_consistency({"designation": "Data Scientist", "skills": ["Python", "SQL", "Statistics"]})
        self.assertGreaterEqual(res["consistency_score"], 50)

    def test_profession_5_doctor(self):
        res = self.consistency_checker.check_consistency({"designation": "Doctor", "skills": ["Patient Care", "Diagnosis"]})
        self.assertGreaterEqual(res["consistency_score"], 50)

    def test_profession_6_lawyer(self):
        res = self.consistency_checker.check_consistency({"designation": "Lawyer", "skills": ["Litigation", "Contracts"]})
        self.assertGreaterEqual(res["consistency_score"], 50)

    def test_profession_7_designer(self):
        res = self.consistency_checker.check_consistency({"designation": "UI/UX Designer", "skills": ["Figma", "Photoshop"]})
        self.assertGreaterEqual(res["consistency_score"], 50)

    def test_profession_8_teacher(self):
        res = self.consistency_checker.check_consistency({"designation": "Teacher", "skills": ["Curriculum Development", "Pedagogy"]})
        self.assertGreaterEqual(res["consistency_score"], 50)

    def test_profession_9_civil_engineer(self):
        res = self.consistency_checker.check_consistency({"designation": "Civil Engineer", "skills": ["AutoCAD", "Structural Analysis"]})
        self.assertGreaterEqual(res["consistency_score"], 50)

    def test_profession_10_mechanical_engineer(self):
        res = self.consistency_checker.check_consistency({"designation": "Mechanical Engineer", "skills": ["CAD", "SolidWorks"]})
        self.assertGreaterEqual(res["consistency_score"], 50)

    def test_profession_11_chemical_engineer(self):
        res = self.consistency_checker.check_consistency({"designation": "Chemical Engineer", "skills": ["Process Engineering"]})
        self.assertGreaterEqual(res["consistency_score"], 50)

    def test_profession_12_accountant(self):
        res = self.consistency_checker.check_consistency({"designation": "Accountant", "skills": ["Accounting", "Taxation", "Excel"]})
        self.assertGreaterEqual(res["consistency_score"], 50)

    def test_profession_13_hr_professional(self):
        res = self.consistency_checker.check_consistency({"designation": "HR Manager", "skills": ["Recruitment", "Employee Relations"]})
        self.assertGreaterEqual(res["consistency_score"], 50)

    def test_profession_14_marketing_professional(self):
        res = self.consistency_checker.check_consistency({"designation": "Marketing Lead", "skills": ["SEO", "Digital Marketing"]})
        self.assertGreaterEqual(res["consistency_score"], 50)


class SourceTrackerTestCase(APITestCase):
    """
    Test suite for Stage 9 / Phase 9.7 Source Tracking Engine, covering:
    SourceTracker, AuditEngine, HistoryManager, VersionTracker, ChangeLogger,
    ProvenanceEngine, API endpoints, and 10 distinct resume genres.
    """

    def setUp(self):
        self.user = User.objects.create_user(
            username="source_user",
            email="source@example.com",
            password="Password123!"
        )
        self.client.force_authenticate(user=self.user)

        from apps.resumes.services.source_tracker import SourceTracker
        from apps.resumes.services.audit_engine import AuditEngine
        from apps.resumes.services.history_manager import HistoryManager
        from apps.resumes.services.version_tracker import VersionTracker
        from apps.resumes.services.change_logger import ChangeLogger
        from apps.resumes.services.provenance_engine import ProvenanceEngine

        self.source_tracker = SourceTracker()
        self.audit_engine = AuditEngine()
        self.history_manager = HistoryManager()
        self.version_tracker = VersionTracker()
        self.change_logger = ChangeLogger()
        self.provenance_engine = ProvenanceEngine()

    def test_source_tracker_color_mapping(self):
        self.assertEqual(self.source_tracker.get_ui_color("regex"), "#3B82F6")
        self.assertEqual(self.source_tracker.get_ui_color("spacy"), "#8B5CF6")
        self.assertEqual(self.source_tracker.get_ui_color("gemini"), "#F97316")
        self.assertEqual(self.source_tracker.get_ui_color("recovery_engine"), "#10B981")
        self.assertEqual(self.source_tracker.get_ui_color("user_edit"), "#06B6D4")

    def test_audit_engine_explainability(self):
        prov = {"value": "MIT", "source": "recovery_engine", "confidence": 95.0, "reason": "Moved from Skills"}
        explanation = self.audit_engine.generate_field_explanation("education", prov)
        self.assertIn("MIT", explanation["explanation"])
        self.assertEqual(explanation["primary_source"], "recovery_engine")

    def test_version_tracker_diff(self):
        old = {"skills": ["Python"]}
        new = {"skills": ["Python", "Docker"]}
        diff = self.version_tracker.compute_json_diff(old, new)
        self.assertTrue(diff["has_changes"])
        self.assertIn("skills", diff["modified_fields"])

    def test_provenance_engine_full_flow(self):
        payload = {"name": "Alex Taylor", "email": "alex@example.com", "skills": ["Python"]}
        res = self.provenance_engine.process_provenance(payload)
        self.assertIn("provenance_map", res)
        self.assertIn("audit_summary", res)
        self.assertEqual(res["metrics"]["traceability"], 100.0)

    def test_source_api_post_raw_payload(self):
        response = self.client.post("/api/source/", {
            "payload": {"name": "Jane Doe", "email": "jane@example.com"}
        }, format="json")
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        data = response.json()
        self.assertIn("provenance_map", data)

    def test_source_api_post_resume_id(self):
        resume = Resume.objects.create(
            user=self.user,
            resume_title="Source Test Resume",
            master_resume_json={"name": "Jane Doe", "email": "jane@example.com"}
        )
        response = self.client.post("/api/source/", {
            "resume_id": str(resume.id)
        }, format="json")
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        data = response.json()
        self.assertEqual(data["resume_id"], str(resume.id))
        self.assertTrue(FieldSource.objects.filter(resume=resume).exists())

    def test_source_api_history(self):
        resume = Resume.objects.create(user=self.user, resume_title="History Test")
        FieldHistory.objects.create(
            resume=resume,
            field="skills",
            previous_value="Python",
            current_value="Python, Docker",
            source="User Edit"
        )
        response = self.client.get("/api/source/history/")
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        data = response.json()
        self.assertGreaterEqual(data["history_count"], 1)

    def test_source_api_audit(self):
        Resume.objects.create(
            user=self.user,
            resume_title="Audit Test Resume",
            master_resume_json={"name": "Alex Taylor"}
        )
        response = self.client.get("/api/source/audit/")
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        data = response.json()
        self.assertGreaterEqual(data["audited_resumes_count"], 1)

    # 10 GENRES COVERAGE
    def test_genre_1_software(self):
        res = self.provenance_engine.process_provenance({"skills": ["Python", "Docker"]})
        self.assertEqual(res["provenance_map"]["skills"]["source"], "gemini")

    def test_genre_2_student(self):
        res = self.provenance_engine.process_provenance({"name": "Student Alex", "email": "alex@uni.edu"})
        self.assertEqual(res["provenance_map"]["email"]["source"], "regex")

    def test_genre_3_academic_cv(self):
        res = self.provenance_engine.process_provenance({"education": "PhD Harvard University"})
        self.assertIn("education", res["provenance_map"])

    def test_genre_4_designer(self):
        res = self.provenance_engine.process_provenance({"skills": ["Figma", "UI/UX"]})
        self.assertIn("skills", res["provenance_map"])

    def test_genre_5_medical(self):
        res = self.provenance_engine.process_provenance({"designation": "Medical Doctor"})
        self.assertIn("designation", res["provenance_map"])

    def test_genre_6_research(self):
        res = self.provenance_engine.process_provenance({"publications": ["Nature Paper 2024"]})
        self.assertIn("publications", res["provenance_map"])

    def test_genre_7_law(self):
        res = self.provenance_engine.process_provenance({"designation": "Attorney At Law"})
        self.assertIn("designation", res["provenance_map"])

    def test_genre_8_canva(self):
        res = self.provenance_engine.process_provenance({"name": "Canva User", "phone": "555-0199"})
        self.assertEqual(res["provenance_map"]["phone"]["source"], "regex")

    def test_genre_9_broken(self):
        res = self.provenance_engine.process_provenance({})
        self.assertEqual(res["metrics"]["traceability"], 100.0)

    def test_genre_10_multipage(self):
        res = self.provenance_engine.process_provenance({"name": "Multi Page", "experience": ["Job 1", "Job 2"]})
        self.assertIn("experience", res["provenance_map"])


class SelfHealingParserTestCase(APITestCase):
    """
    Test suite for Stage 9 / Phase 9.8 Self-Healing Parser, covering:
    ResultMerger, DecisionEngine, PipelineManager, PipelineOrchestrator,
    MasterResumeBuilder, SelfHealingParser orchestrator, API endpoints,
    and 14 distinct resume genres.
    """

    def setUp(self):
        self.user = User.objects.create_user(
            username="healing_user",
            email="healing@example.com",
            password="Password123!"
        )
        self.client.force_authenticate(user=self.user)

        from apps.resumes.services.result_merger import ResultMerger
        from apps.resumes.services.decision_engine import DecisionEngine
        from apps.resumes.services.pipeline_manager import PipelineManager
        from apps.resumes.services.orchestrator import PipelineOrchestrator
        from apps.resumes.services.resume_builder import MasterResumeBuilder
        from apps.resumes.services.self_healing_parser import SelfHealingParser

        self.merger = ResultMerger()
        self.decision_engine = DecisionEngine()
        self.pipeline_manager = PipelineManager()
        self.orchestrator = PipelineOrchestrator()
        self.builder = MasterResumeBuilder()
        self.parser = SelfHealingParser()

    def test_result_merger_confidence_resolution(self):
        merged = self.merger.merge_stage_results(
            regex_data={"email": "regex@example.com"},
            spacy_data={"email": "spacy@example.com"},
            gemini_data={"email": "gemini@example.com"}
        )
        self.assertEqual(merged["email"], "regex@example.com")

    def test_decision_engine_auto_approval_tiers(self):
        d1 = self.decision_engine.evaluate_decision(96.0, 5, 5, 5)
        self.assertEqual(d1["decision"], "accept")
        self.assertEqual(d1["approval_tier"], "auto_approve")

        d2 = self.decision_engine.evaluate_decision(88.0, 5, 4, 4)
        self.assertEqual(d2["decision"], "review")

    def test_master_resume_builder_structure(self):
        res = self.builder.build_master_resume(
            merged_payload={"name": "Alex Taylor", "email": "alex@example.com", "skills": ["Python"]},
            confidence_score=96.0
        )
        self.assertIn("profile", res)
        self.assertEqual(res["profile"]["name"], "Alex Taylor")
        self.assertIn("metadata", res)
        self.assertEqual(res["metadata"]["quality_scores"]["extraction_accuracy"], 98.0)

    def test_self_healing_parser_full_flow(self):
        payload = {"name": "Alex Taylor", "email": "alex@example.com", "skills": ["Python"]}
        res = self.parser.parse_and_heal(payload)
        self.assertIn("master_resume", res)
        self.assertIn("healing_report", res)
        self.assertEqual(res["healing_report"]["decision"], "accept")

    def test_healing_api_post_raw_payload(self):
        response = self.client.post("/api/self-healing/", {
            "payload": {"name": "Jane Doe", "email": "jane@example.com"}
        }, format="json")
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        data = response.json()
        self.assertIn("master_resume", data)

    def test_healing_api_post_resume_id(self):
        resume = Resume.objects.create(
            user=self.user,
            resume_title="Healing Test Resume",
            master_resume_json={"name": "Jane Doe", "email": "jane@example.com"}
        )
        response = self.client.post("/api/self-healing/", {
            "resume_id": str(resume.id)
        }, format="json")
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        data = response.json()
        self.assertEqual(data["resume_id"], str(resume.id))
        self.assertTrue(SelfHealingReport.objects.filter(resume=resume).exists())

    def test_healing_api_report_list(self):
        resume = Resume.objects.create(user=self.user, resume_title="Report Test")
        SelfHealingReport.objects.create(
            resume=resume,
            confidence=96.0,
            decision="accept",
            summary="All clean"
        )
        response = self.client.get("/api/self-healing/report/")
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        data = response.json()
        self.assertGreaterEqual(data["total_reports"], 1)

    # 14 GENRES COVERAGE
    def test_genre_1_software(self):
        res = self.parser.parse_and_heal({"name": "Dev", "skills": ["Python", "Docker"]})
        self.assertEqual(res["master_resume"]["profile"]["name"], "Dev")

    def test_genre_2_student(self):
        res = self.parser.parse_and_heal({"name": "Student", "education": [{"degree": "BS"}]})
        self.assertTrue(len(res["master_resume"]["education"]) > 0)

    def test_genre_3_designer(self):
        res = self.parser.parse_and_heal({"skills": ["Figma", "Photoshop"]})
        self.assertIn("skills", res["master_resume"])

    def test_genre_4_medical(self):
        res = self.parser.parse_and_heal({"designation": "Doctor", "skills": ["Patient Care"]})
        self.assertEqual(res["master_resume"]["profile"]["designation"], "Doctor")

    def test_genre_5_research(self):
        res = self.parser.parse_and_heal({"projects": ["Machine Learning Research"]})
        self.assertTrue(len(res["master_resume"]["projects"]) > 0)

    def test_genre_6_teacher(self):
        res = self.parser.parse_and_heal({"designation": "Teacher"})
        self.assertEqual(res["master_resume"]["profile"]["designation"], "Teacher")

    def test_genre_7_academic_cv(self):
        res = self.parser.parse_and_heal({"education": [{"degree": "PhD"}]})
        self.assertTrue(len(res["master_resume"]["education"]) > 0)

    def test_genre_8_law(self):
        res = self.parser.parse_and_heal({"designation": "Lawyer"})
        self.assertEqual(res["master_resume"]["profile"]["designation"], "Lawyer")

    def test_genre_9_twocolumn(self):
        res = self.parser.parse_and_heal({"name": "Two Column", "email": "tc@example.com"})
        self.assertEqual(res["master_resume"]["profile"]["email"], "tc@example.com")

    def test_genre_10_canva(self):
        res = self.parser.parse_and_heal({"name": "Canva User", "phone": "555-0199"})
        self.assertEqual(res["master_resume"]["profile"]["phone"], "555-0199")

    def test_genre_11_broken_pdf(self):
        res = self.parser.parse_and_heal({})
        self.assertGreaterEqual(res["master_resume"]["metadata"]["confidence"], 70.0)

    def test_genre_12_multipage(self):
        res = self.parser.parse_and_heal({"experience": ["Job A", "Job B"]})
        self.assertTrue(len(res["master_resume"]["experience"]) > 0)

    def test_genre_13_large(self):
        res = self.parser.parse_and_heal({"skills": ["Skill " + str(i) for i in range(50)]})
        self.assertEqual(len(res["master_resume"]["skills"]), 50)

    def test_genre_14_malformed(self):
        res = self.parser.parse_and_heal({"skills": ["C++", "c++"]})
        self.assertIn("master_resume", res)


class ResumeCopilotTestCase(APITestCase):
    """
    Test suite for Stage 9 / Phase 9.9 Resume Copilot, covering:
    ConversationEngine, ResumeEditor, ATSAdvisor, SummaryGenerator,
    ChangeManager, MemoryEngine, SuggestionEngine, ProfileOptimizer,
    ResumeCopilot orchestrator, API endpoints, and 12 distinct resume genres.
    """

    def setUp(self):
        self.user = User.objects.create_user(
            username="copilot_user",
            email="copilot@example.com",
            password="Password123!"
        )
        self.client.force_authenticate(user=self.user)

        from apps.resumes.services.conversation_engine import ConversationEngine
        from apps.resumes.services.resume_editor import ResumeEditor
        from apps.resumes.services.ats_advisor import ATSAdvisor
        from apps.resumes.services.summary_generator import SummaryGenerator
        from apps.resumes.services.change_manager import ChangeManager
        from apps.resumes.services.memory_engine import MemoryEngine
        from apps.resumes.services.suggestion_engine import SuggestionEngine
        from apps.resumes.services.profile_optimizer import ProfileOptimizer
        from apps.resumes.services.resume_copilot import ResumeCopilot

        self.conversation_engine = ConversationEngine()
        self.editor = ResumeEditor()
        self.ats_advisor = ATSAdvisor()
        self.summary_generator = SummaryGenerator()
        self.change_manager = ChangeManager()
        self.memory_engine = MemoryEngine()
        self.suggestion_engine = SuggestionEngine()
        self.optimizer = ProfileOptimizer()
        self.copilot = ResumeCopilot()

    def test_conversation_engine_intent_parsing(self):
        i1 = self.conversation_engine.parse_user_intent("Add Docker")
        self.assertEqual(i1["intent"], "add_skill")
        self.assertEqual(i1["target"], "Docker")

        i2 = self.conversation_engine.parse_user_intent("Remove Java")
        self.assertEqual(i2["intent"], "remove_skill")
        self.assertEqual(i2["target"], "Java")

        i3 = self.conversation_engine.parse_user_intent("Improve my ATS")
        self.assertEqual(i3["intent"], "improve_ats")

    def test_resume_editor_mutations(self):
        base = {"skills": ["Python"]}
        added = self.editor.add_skill(base, "Docker")
        self.assertIn("Docker", added["skills"])

        removed = self.editor.remove_skill(added, "Python")
        self.assertNotIn("Python", removed["skills"])

    def test_ats_advisor_predictions(self):
        res = self.ats_advisor.analyze_ats({"skills": ["Python"]})
        self.assertIn("current_ats", res)
        self.assertIn("estimated_ats", res)
        self.assertGreater(res["estimated_ats"], res["current_ats"])

    def test_memory_engine_preferences(self):
        base = {"skills": ["ReactJS"], "education": [{"institution": "LJ Institute"}]}
        pref = self.memory_engine.apply_user_preferences(base)
        self.assertIn("React", pref["skills"])
        self.assertEqual(pref["education"][0]["institution"], "LJ University")

    def test_copilot_full_chat_flow(self):
        payload = {"skills": ["Python"]}
        res = self.copilot.process_chat("Add Docker", payload)
        self.assertEqual(res["intent"], "add_skill")
        self.assertIn("Docker", res["updated_master_json"]["skills"])

    def test_copilot_api_chat(self):
        resume = Resume.objects.create(
            user=self.user,
            resume_title="Copilot Test Resume",
            master_resume_json={"skills": ["Python"]}
        )
        response = self.client.post("/api/copilot/chat/", {
            "resume_id": str(resume.id),
            "message": "Add Docker"
        }, format="json")
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        data = response.json()
        self.assertIn("Docker", data["updated_master_json"]["skills"])
        self.assertTrue(CopilotConversation.objects.filter(user=self.user).exists())
        self.assertTrue(CopilotAction.objects.filter(resume=resume).exists())

    def test_copilot_api_undo_redo(self):
        resume = Resume.objects.create(
            user=self.user,
            resume_title="Undo Test Resume",
            master_resume_json={"skills": ["Python"]}
        )
        self.client.post("/api/copilot/chat/", {"resume_id": str(resume.id), "message": "Add Docker"}, format="json")
        
        # Test Undo
        undo_res = self.client.post("/api/copilot/action/", {"resume_id": str(resume.id), "action": "undo"}, format="json")
        self.assertEqual(undo_res.status_code, status.HTTP_200_OK)
        self.assertNotIn("Docker", undo_res.json()["master_resume_json"]["skills"])

        # Test Redo
        redo_res = self.client.post("/api/copilot/action/", {"resume_id": str(resume.id), "action": "redo"}, format="json")
        self.assertEqual(redo_res.status_code, status.HTTP_200_OK)
        self.assertIn("Docker", redo_res.json()["master_resume_json"]["skills"])

    def test_copilot_api_history_and_suggestions(self):
        Resume.objects.create(user=self.user, resume_title="Sugg Resume", master_resume_json={})
        h_res = self.client.get("/api/copilot/history/")
        self.assertEqual(h_res.status_code, status.HTTP_200_OK)

        s_res = self.client.get("/api/copilot/suggestions/")
        self.assertEqual(s_res.status_code, status.HTTP_200_OK)

        c_res = self.client.get("/api/copilot/changes/")
        self.assertEqual(c_res.status_code, status.HTTP_200_OK)

    # 12 GENRES COVERAGE
    def test_genre_1_software(self):
        res = self.copilot.process_chat("Improve my ATS", {"skills": ["Python", "Django"]})
        self.assertEqual(res["intent"], "improve_ats")

    def test_genre_2_designer(self):
        res = self.copilot.process_chat("Add Figma", {"skills": ["Photoshop"]})
        self.assertIn("Figma", res["updated_master_json"]["skills"])

    def test_genre_3_student(self):
        res = self.copilot.process_chat("My education is wrong", {"education": []})
        self.assertEqual(res["intent"], "fix_education")

    def test_genre_4_researcher(self):
        res = self.copilot.process_chat("Generate stronger summary", {"profile": {"designation": "Researcher"}})
        self.assertEqual(res["intent"], "generate_summary")

    def test_genre_5_teacher(self):
        res = self.copilot.process_chat("Add Curriculum Design", {"skills": ["Teaching"]})
        self.assertIn("Curriculum Design", res["updated_master_json"]["skills"])

    def test_genre_6_doctor(self):
        res = self.copilot.process_chat("Improve summary", {"profile": {"designation": "Doctor"}})
        self.assertEqual(res["intent"], "generate_summary")

    def test_genre_7_lawyer(self):
        res = self.copilot.process_chat("Why is MIT inside Education?", {})
        self.assertEqual(res["intent"], "explain_parser")

    def test_genre_8_fresher(self):
        res = self.copilot.process_chat("Improve my ATS", {})
        self.assertIn("ats_summary", res)

    def test_genre_9_academic_cv(self):
        res = self.copilot.process_chat("Generate summary", {"profile": {"designation": "Professor"}})
        self.assertIn("Professor", res["response"])

    def test_genre_10_twocolumn(self):
        res = self.copilot.process_chat("Add Python", {})
        self.assertIn("Python", res["updated_master_json"]["skills"])

    def test_genre_11_canva(self):
        res = self.copilot.process_chat("Remove Java", {"skills": ["Java", "C++"]})
        self.assertNotIn("Java", res["updated_master_json"]["skills"])

    def test_genre_12_broken(self):
        res = self.copilot.process_chat("General question", {})
        self.assertEqual(res["intent"], "chat")













