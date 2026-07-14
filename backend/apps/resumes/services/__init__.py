import os
import time
import logging
import unicodedata
from django.utils import timezone
import fitz  # PyMuPDF
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table

from ..models import Resume

logger = logging.getLogger(__name__)


class ResumeExtractionService:
    """
    Service for extracting clean text from uploaded resumes (PDF & DOCX).
    Handles file detection, text extraction, normalization, and updating Resume model.
    """

    def extract_resume_text(self, resume: Resume) -> bool:
        """
        Executes the text extraction flow for a given Resume instance.
        Updates model fields and returns True on success, False on failure.
        """
        start_time = time.time()
        resume.extraction_status = Resume.ExtractionStatus.PROCESSING
        resume.error_message = ""
        resume.save(update_fields=["extraction_status", "error_message"])

        logger.info(f"Start extraction for resume '{resume.id}' (type: {resume.mime_type})")

        # Basic validation
        if resume.is_deleted:
            self._handle_failure(resume, "Cannot extract text from a deleted resume.")
            return False

        if not resume.original_file:
            self._handle_failure(resume, "Resume file is missing on storage.")
            return False

        file_path = resume.original_file.path
        if not os.path.exists(file_path):
            self._handle_failure(resume, f"Resume file path does not exist on disk: {file_path}")
            return False

        ext = os.path.splitext(resume.original_filename)[1].lower()
        if not ext:
            ext = os.path.splitext(resume.original_file.name)[1].lower()

        try:
            # Route based on file extension
            if ext == ".pdf":
                raw_text = self._extract_pdf(file_path)
            elif ext == ".docx":
                raw_text = self._extract_docx(file_path)
            else:
                raise ValueError(f"Unsupported file extension '{ext}'. Only PDF and DOCX are allowed.")

            # Normalization
            cleaned_text = self._normalize_text(raw_text)

            # Check if extraction yielded any text
            if not cleaned_text.strip():
                raise ValueError("No readable text could be extracted from this document.")

            # Update Resume model on success
            duration = time.time() - start_time
            resume.extracted_text = cleaned_text
            resume.extraction_status = Resume.ExtractionStatus.COMPLETED
            resume.extraction_time = timezone.now()
            resume.processing_duration = round(duration, 3)
            resume.error_message = ""
            resume.save(
                update_fields=[
                    "extracted_text",
                    "extraction_status",
                    "extraction_time",
                    "processing_duration",
                    "error_message",
                ]
            )

            logger.info(
                f"Successfully completed extraction for resume '{resume.id}' in {duration:.3f}s. "
                f"Text length: {len(cleaned_text)} chars."
            )
            return True

        except Exception as e:
            duration = time.time() - start_time
            error_msg = str(e)
            logger.error(f"Failed extraction for resume '{resume.id}': {error_msg}", exc_info=True)
            self._handle_failure(resume, error_msg, duration=duration)
            return False

    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from a PDF file using PyMuPDF (fitz)."""
        text_parts = []
        try:
            doc = fitz.open(file_path)
        except Exception as e:
            raise ValueError(f"The PDF file is corrupted or could not be opened: {str(e)}")

        if doc.is_encrypted:
            raise ValueError("The PDF file is password-protected or encrypted.")

        pages_processed = 0
        for page_num in range(len(doc)):
            # Cap execution safety check for extremely large documents
            if page_num >= 100:
                logger.warning("PDF exceeds 100 pages. Truncating extraction for performance.")
                break
            
            page = doc.load_page(page_num)
            page_text = page.get_text("text")
            if page_text.strip():
                text_parts.append(page_text)
                pages_processed += 1

        doc.close()
        logger.info(f"Processed {pages_processed} pages from PDF.")
        return "\n\n".join(text_parts)

    def _extract_docx(self, file_path: str) -> str:
        """Extract text, tables, headers, and footers from a DOCX file using python-docx."""
        try:
            doc = Document(file_path)
        except Exception as e:
            raise ValueError(f"The DOCX file is corrupted or could not be opened: {str(e)}")

        text_parts = []

        # 1. Extract headers (if any)
        for section in doc.sections:
            if section.header and not section.header.is_linked_to_previous:
                for p in section.header.paragraphs:
                    p_text = p.text.strip()
                    if p_text:
                        text_parts.append(p_text)

        # 2. Extract body paragraphs & tables in correct reading order
        for element in doc.element.body:
            if element.tag.endswith("p"):
                p = Paragraph(element, doc)
                p_text = p.text.strip()
                if p_text:
                    text_parts.append(p_text)
            elif element.tag.endswith("tbl"):
                table = Table(element, doc)
                table_text_parts = []
                for row in table.rows:
                    row_cells_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    # De-duplicate adjacent identical cells due to cell merging
                    unique_cells = []
                    for cell_text in row_cells_text:
                        if not unique_cells or unique_cells[-1] != cell_text:
                            unique_cells.append(cell_text)
                    if unique_cells:
                        table_text_parts.append(" | ".join(unique_cells))
                if table_text_parts:
                    text_parts.append("\n".join(table_text_parts))

        # 3. Extract footers (if any)
        for section in doc.sections:
            if section.footer and not section.footer.is_linked_to_previous:
                for p in section.footer.paragraphs:
                    p_text = p.text.strip()
                    if p_text:
                        text_parts.append(p_text)

        return "\n\n".join(text_parts)

    def _normalize_text(self, text: str) -> str:
        """Normalize unicode normalization form, control chars, white space, and double lines."""
        if not text:
            return ""

        # Normalize unicode characters to NFKC form (standardizes special symbols, ligatures, etc.)
        text = unicodedata.normalize("NFKC", text)

        # Filter out control characters except tabs, newlines, and carriage returns
        text = "".join(
            ch for ch in text 
            if ch in ("\n", "\r", "\t") or unicodedata.category(ch)[0] != "C"
        )

        # Standardize line endings to LF (\n)
        text = text.replace("\r\n", "\n").replace("\r", "\n")

        # Convert tabs to 4 spaces
        text = text.replace("\t", "    ")

        # Strip trailing space on each line
        lines = [line.rstrip() for line in text.split("\n")]

        # Collapse multiple blank lines down to a maximum of 1 blank line (max 2 consecutive newlines)
        normalized_lines = []
        blank_count = 0
        for line in lines:
            if not line:
                blank_count += 1
                if blank_count <= 1:
                    normalized_lines.append("")
            else:
                blank_count = 0
                normalized_lines.append(line)

        text = "\n".join(normalized_lines)
        return text.strip()

    def _handle_failure(self, resume: Resume, error_message: str, duration: float = None):
        """Update the model status to failed with details."""
        resume.extraction_status = Resume.ExtractionStatus.FAILED
        resume.error_message = error_message
        if duration is not None:
            resume.processing_duration = round(duration, 3)
        resume.save(update_fields=["extraction_status", "error_message", "processing_duration"])


from .knowledge_base import KnowledgeBase
from .ontology_engine import OntologyEngine
from .entity_classifier import EntityClassifier
from .semantic_matcher import SemanticMatcher
from .semantic_validator import SemanticValidator
from .duplicate_detector import DuplicateDetector
from .timeline_validator import TimelineValidator
from .contact_validator import ContactValidator
from .consistency_validator import ConsistencyValidator
from .quality_validator import QualityValidator
from .section_error_detector import SectionErrorDetector
from .error_detector import ErrorDetector
from .entity_mover import EntityMover
from .duplicate_resolver import DuplicateResolver
from .date_recovery import DateRecovery
from .summary_recovery import SummaryRecovery
from .recovery_rules import RecoveryRules
from .recovery_engine import RecoveryEngine, RecoveryPlanner
from .timeline_checker import TimelineChecker
from .career_checker import CareerChecker
from .role_checker import RoleChecker
from .profile_checker import ProfileChecker
from .completeness_checker import CompletenessChecker
from .consistency_checker import ConsistencyChecker
from .source_tracker import SourceTracker
from .audit_engine import AuditEngine
from .history_manager import HistoryManager
from .version_tracker import VersionTracker
from .change_logger import ChangeLogger
from .provenance_engine import ProvenanceEngine
from .result_merger import ResultMerger
from .decision_engine import DecisionEngine
from .pipeline_manager import PipelineManager
from .orchestrator import PipelineOrchestrator
from .resume_builder import MasterResumeBuilder
from .self_healing_parser import SelfHealingParser
from .conversation_engine import ConversationEngine
from .resume_editor import ResumeEditor
from .ats_advisor import ATSAdvisor
from .summary_generator import SummaryGenerator
from .change_manager import ChangeManager
from .memory_engine import MemoryEngine
from .suggestion_engine import SuggestionEngine
from .profile_optimizer import ProfileOptimizer
from .resume_copilot import ResumeCopilot

__all__ = [
    "ResumeExtractionService",
    "KnowledgeBase",
    "OntologyEngine",
    "EntityClassifier",
    "SemanticMatcher",
    "SemanticValidator",
    "DuplicateDetector",
    "TimelineValidator",
    "ContactValidator",
    "ConsistencyValidator",
    "QualityValidator",
    "SectionErrorDetector",
    "ErrorDetector",
    "EntityMover",
    "DuplicateResolver",
    "DateRecovery",
    "SummaryRecovery",
    "RecoveryRules",
    "RecoveryEngine",
    "RecoveryPlanner",
    "TimelineChecker",
    "CareerChecker",
    "RoleChecker",
    "ProfileChecker",
    "CompletenessChecker",
    "ConsistencyChecker",
    "SourceTracker",
    "AuditEngine",
    "HistoryManager",
    "VersionTracker",
    "ChangeLogger",
    "ProvenanceEngine",
    "ResultMerger",
    "DecisionEngine",
    "PipelineManager",
    "PipelineOrchestrator",
    "MasterResumeBuilder",
    "SelfHealingParser",
    "ConversationEngine",
    "ResumeEditor",
    "ATSAdvisor",
    "SummaryGenerator",
    "ChangeManager",
    "MemoryEngine",
    "SuggestionEngine",
    "ProfileOptimizer",
    "ResumeCopilot",
]

